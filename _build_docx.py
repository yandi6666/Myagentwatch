"""
Build MyAgentWatch project introduction .docx files in the v7 style.
Chinese version + English version, updated to v6.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path
import datetime

PROJECT_DIR = Path(__file__).resolve().parent
TODAY = datetime.date.today().strftime("%Y-%m-%d")

# ── color palette ──────────────────────────────────────────────
BG_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
ACCENT    = RGBColor(0x0F, 0x34, 0x60)
HIGHLIGHT = RGBColor(0xE9, 0x45, 0x60)
GRAY      = RGBColor(0x66, 0x66, 0x66)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
BLACK     = RGBColor(0x00, 0x00, 0x00)
TABLE_HDR = RGBColor(0x0F, 0x34, 0x60)
TABLE_BG  = RGBColor(0xF5, 0xF7, 0xFA)

# ── helpers ─────────────────────────────────────────────────────

def hdr(doc, text, level=1):
    """Add a heading with v7 style."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run.font.color.rgb = BG_DARK if level == 1 else ACCENT
    return h

def para(doc, text, bold=False, italic=False, size=10.5, color=None, align=None, spacing_after=6, indent=None):
    """Add a paragraph with v7 style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.35
    if align is not None:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def bullet(doc, text, level=0, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    p.clear()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = 'Microsoft YaHei'
        r.font.size = Pt(10.5)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10.5)
    return p

def section_label(doc, text):
    """Add a ▎-style section label within a chapter."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    run = p.add_run(f"▎ {text}")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = ACCENT
    return p

def make_table(doc, headers, rows, col_widths=None):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F3460" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9)
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA" w:val="clear"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for ri, row in enumerate(table.rows):
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()  # spacer
    return table

def page_break(doc):
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════
# CHINESE VERSION
# ═══════════════════════════════════════════════════════════════════

def build_cn():
    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.35

    # Page margins
    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MyAgentWatch")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = BG_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("多 Agent 实时监控与协作平台")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("打破 AI Agent 黑盒 —— 实时追踪全链路行为与成本，安全审批每一步执行")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"版本 v6    |    {TODAY}\nPython 3.12  ·  Flask  ·  SocketIO  ·  SQLite  ·  Chart.js\n许可证：AGPL-3.0-only")
    run.font.name = 'Microsoft YaHei'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    page_break(doc)

    # ── TABLE OF CONTENTS (manual) ──────────────────────────
    hdr(doc, "目录", level=1)
    toc_items = [
        "一、项目背景：为什么需要 MyAgentWatch？",
        "    1.1 行业趋势：AI Agent 正在成为开发团队的标准配置",
        "    1.2 三大核心痛点",
        "    1.3 MyAgentWatch 的答案",
        "二、项目概述",
        "    2.1 MyAgentWatch 是什么？",
        "    2.2 核心价值主张",
        "    2.3 关键数据",
        "    2.4 适用场景",
        "三、Before / After：使用前后的差异",
        "四、v6 核心功能详解",
        "    4.1 网页群聊 —— 协作主入口",
        "    4.2 安全执行闭环 —— 层层设防",
        "    4.3 Agent Task Queue —— 结构化任务系统",
        "    4.4 SSE 实时事件流 —— Agent 的「直播画面」",
        "    4.5 Token 用量仪表盘 —— 省钱的第一道防线",
        "    4.6 CLI / daemon —— Agent 端工具",
        "    4.7 告警规则引擎 —— 主动发现问题",
        "五、技术架构",
        "    5.1 四层架构模型",
        "    5.2 适配器模式",
        "    5.3 Agent 状态机",
        "    5.4 双通道实时推送",
        "    5.5 关键设计决策",
        "六、数据模型设计",
        "七、与 Multica 的竞品对比",
        "八、产品哲学",
        "九、项目状态与路线图",
        "十、部署与运维",
        "附录：快速参考表",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10)
        if not item.startswith("    "):
            run.bold = True

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 1: WHY
    # ══════════════════════════════════════════════════════════
    hdr(doc, "一、项目背景：为什么需要 MyAgentWatch？", level=1)

    hdr(doc, "1.1 行业趋势：AI Agent 正在成为开发团队的标准配置", level=2)
    para(doc, "2025 年以来，以 Claude Code、OpenCode、Codex、Cursor、Gemini CLI 为代表的 AI 编程 Agent 迅速普及。"
         "越来越多的团队将 Agent 深度集成到日常开发流程中——代码生成、需求分析、架构设计、测试编写、Bug 修复。"
         "在量化交易等高频决策场景，AI Agent 还承担策略研究、因子挖掘、回测分析等关键任务。"
         "Agent 已经从「辅助工具」演变为「虚拟团队成员」，其运行状态、决策质量和资源消耗直接影响项目进度和成本。")

    hdr(doc, "1.2 三大核心痛点", level=2)

    para(doc, "痛点一：Agent 间通信是黑盒——你看不到 Agent 之间说了什么", bold=True, size=11)
    para(doc, "以一个典型的多 Agent 协作场景为例（假设场景，用于说明问题）：OpenClaw 是一款通用 AI Agent 编排框架（类似 AI 团队的「项目经理」），"
         "OpenCode 是一款专注于终端编程的 AI Agent（类似 AI 团队的「程序员」）。典型工作流是：用户给 OpenClaw 下达任务 → "
         "OpenClaw 分析后将编码部分委托给 OpenCode → OpenCode 完成编码把结果返回给 OpenClaw → OpenClaw 汇总呈现给用户。")

    # Before
    section_label(doc, "Before（使用 MyAgentWatch 之前）—— 用户视角")
    bullets_before = [
        "用户在 OpenClaw 的聊天界面中说：「帮我写一个 Flask REST API」",
        "OpenClaw 内部通过 ACP（Agent Communication Protocol，JSON-RPC 2.0 over NDJSON 子进程协议）启动了一个 OpenCode 子进程",
        "OpenClaw 向 OpenCode 发送了编码指令，但这是底层 IPC 通信——用户完全看不到指令内容",
        "OpenCode 开始读文件、写代码、跑测试，整个过程在终端子进程中静默执行",
        "OpenCode 完成后返回结果给 OpenClaw，OpenClaw 用自己的语言复述结果给用户",
        "用户的困惑：OpenClaw 到底让 OpenCode 做了什么？指令是否准确？OpenCode 有没有误解？OpenCode 报了什么错？OpenClaw 有没有篡改/简化 OpenCode 的输出？全部不可见",
        "更糟的是：如果 OpenCode 陷入死循环（反复读同一个文件）、调用了错误的工具、或者返回了错误的结果，用户在 OpenClaw 界面完全感知不到——只能看到 OpenClaw 说「处理好了」或「遇到了一些问题」",
    ]
    for b in bullets_before:
        bullet(doc, b)

    # Dead loop cost
    section_label(doc, "OpenCode 陷入死循环的真实代价")
    para(doc, "假设 OpenCode 在编写代码时遇到了一个它不理解的 import 错误。在没有监控的情况下，一个典型的死循环场景是这样的：")
    dead_loop = [
        "第 1 轮：OpenCode 写入了一段有 bug 的代码 → 运行测试 → 失败 → 「让我看看错误信息」→ 读取错误日志",
        "第 2 轮：OpenCode 「我理解了，是 import 路径问题」→ 修改代码 → 运行测试 → 又失败 → 「不对，让我重新看看」→ 再次读取同一个文件",
        "第 3-8 轮：OpenCode 反复读同一个文件 → 反复猜测原因 → 反复修改 → 反复失败。每次循环消耗 8,000-15,000 input tokens + 2,000-5,000 output tokens",
        "第 9 轮：OpenCode 的思考开始退化（上下文窗口被无效信息填满）→ 「我建议回滚所有改动重来」→ 又开始新一轮循环",
        "最终结果：30 分钟后用户打开 OpenClaw，看到一条轻描淡写的「任务处理遇到一些困难，正在重试」。但实际上：已消耗 380,000 input tokens（成本约 $1.14）、95,000 output tokens（成本约 $0.57）、总共浪费 $1.71、耗时 30 分钟、源文件被反复修改了 12 次、git diff 显示大量无效改动、用户完全不知道这 30 分钟里发生了什么",
        "如果这个场景发生在凌晨 3 点的自动化任务中，团队第二天早上才会发现：任务没完成、Token 被刷掉了 $15+、代码仓库多了一堆需要回滚的提交",
    ]
    for d in dead_loop:
        bullet(doc, d)

    para(doc, "这不是假设。在实际使用中，Claude Opus 4 单次请求可消耗高达 32,000 output tokens（约 $0.48），"
         "如果不加监控，一个循环 10 轮的 Agent 足以在一小时内烧掉 $5-15。"
         "对于同时运行 5 个 Agent 的量化交易团队，一天的死循环损失可能高达 $200。")

    # With MyAgentWatch
    section_label(doc, "而有了 MyAgentWatch")
    bullets_after = [
        "当 OpenCode 第 3 次读取同一个文件时，告警规则触发橙色 blocked 状态",
        "用户实时收到通知：「OpenCode 在过去 5 分钟内读取了同一个文件 3 次，疑似死循环」",
        "用户在群聊中直接 @OpenCode 发送：「停止当前任务，换一种思路」——人工介入，打破循环",
        "或者用户查看 SSE 事件流，发现 OpenCode 的 thinking 内容越来越长、越来越混乱（上下文污染），直接通过 CLI 强制挂起 Agent",
        "Token 仪表盘实时显示 OpenCode 的消耗曲线陡然上升，用户一眼就能发现异常",
    ]
    for b in bullets_after:
        bullet(doc, b)

    para(doc, "这就是透明化的力量：不是阻止 Agent 犯错（Agent 必然会犯错），而是在犯错的第一时间让你知道，并给你干预的手段。")

    para(doc, "这就像两个同事关上门在小房间里讨论，你只听到了最终「我们决定这样」的结论——"
         "中间的争论、纠错、方案变更你全都不知道。对于量化交易等关键场景，这种黑盒是不可接受的："
         "你不知道一个交易策略是经过充分验证的还是草率决定的。")

    # After
    section_label(doc, "After（使用 MyAgentWatch 之后）—— 用户视角")
    bullets_after2 = [
        "用户同样对 OpenClaw 说：「帮我写一个 Flask REST API」",
        "OpenClaw 通过 MyAgentWatch 的群聊系统，在群聊中 @OpenCode 并发送委托消息：「请帮我创建一个 Flask REST API，包含 GET /users 和 POST /users 两个端点，使用 SQLite 存储，代码风格遵循 PEP 8。」",
        "这条委托消息以紫色 Handoff 卡片的形式实时显示在群聊消息流中，清晰标注「OpenClaw → OpenCode / subagent_type: coding / 状态: running」",
        "OpenCode 收到任务后，在群聊中回复自己的执行进展——thinking 消息（蓝色气泡·思考过程）、tool_call 消息（橙色框·工具调用）、response 消息（绿色气泡·代码输出结果）",
        "同时，SSE 事件流实时推送 OpenCode 的每一步操作：思考了 2.3 秒 → 调用了 read 工具读取 app.py → 调用了 write 工具写入 routes/users.py → 调用了 bash 工具运行 pytest → 测试全部通过 → 返回结果给 OpenClaw",
        "OpenClaw 收到结果后，在群聊中用自己的话总结回复用户，同时把 OpenCode 的原始输出以绿色 Share 卡片形式分享到群聊",
        "整个过程完整记录：谁委托了谁、委托了什么、子 Agent 怎么做的、每一步的 Token 消耗和耗时、最终结果是什么——全部可追溯、可审计",
    ]
    for b in bullets_after2:
        bullet(doc, b)

    para(doc, "这就是 MyAgentWatch 群聊系统作为「Agent 间通信总线」的核心价值："
         "让 Agent 之间的每一次对话、每一次委托、每一次交接都透明可见。"
         "团队可以在事后回放任意一次多 Agent 协作的全过程，就像查看 Slack 频道里的讨论记录一样自然。")

    para(doc, "痛点二：Token 成本失控——团队不知道钱花在哪", bold=True, size=11)
    para(doc, "一个中型开发团队每天可能消耗数百万 Token。不同模型价差可达 50 倍以上。"
         "不同 Agent 类型消耗差异巨大（plan 层通常消耗远高于 build 层）。不同时段消耗不均衡。"
         "没有任何工具能告诉团队：「今天花了多少钱？哪个 Agent 最费钱？缓存命中率是多少？有没有模型没用定价？」")

    para(doc, "痛点三：缺乏运维保障——Agent 停了都没人知道", bold=True, size=11)
    para(doc, "Agent 进程可能因为各种原因挂掉：API 限流、Token 耗尽、内存溢出、数据源连接断开。"
         "传统 APM 工具（如 Grafana、Datadog）无法感知 AI Agent 的专属指标：心跳、模型状态、工具调用成功率、缓存效率。"
         "当 Agent 在凌晨 3 点自动化任务中宕掉，团队要等到第二天才发现，浪费大量计算资源和时间窗口。")

    hdr(doc, "1.3 MyAgentWatch 的答案：Agent 通信基础设施", level=2)
    para(doc, "MyAgentWatch 不仅是一个监控仪表盘，更是一套 Agent 之间的通信基础设施。"
         "它的群聊系统实际上充当了 Agent 间通信的统一消息总线：")
    infra_points = [
        ("沟通语言可指定 —— ", "Agent 之间的所有通信必须通过群聊消息，用户可以要求 Agent 使用中文或英文进行对话。"
         "这让非技术团队成员也能看懂 Agent 之间在讨论什么。"),
        ("消息类型区分 —— ", "普通消息（灰色气泡）、工具调用（橙色边框）、Agent 交接（紫色边框）、成果分享（绿色边框），"
         "不同颜色和边框样式让用户一眼就能分辨消息性质，快速定位关键信息。"),
        ("完全可追溯 —— ", "所有 Agent 间通信持久化存储，天然键 natural_key 去重保证数据完整。"
         "用户可以按 trace_id 回放任意一次多 Agent 协作的完整链路。"),
        ("Token 成本归属 —— ", "每条委托消息关联具体的 turn_id，Token 消耗精确归属到发起任务的 Agent。"
         "你可以明确知道「OpenClaw 委托 OpenCode 执行的任务，消耗了 45,000 input tokens + 12,000 output tokens，成本 $0.38」。"),
    ]
    for prefix, text in infra_points:
        bullet(doc, text, bold_prefix=prefix)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 2: OVERVIEW
    # ══════════════════════════════════════════════════════════
    hdr(doc, "二、项目概述", level=1)

    hdr(doc, "2.1 MyAgentWatch 是什么？", level=2)
    para(doc, "MyAgentWatch 是一款开源的、自托管的 AI Agent 实时监控与协作平台。"
         "它以适配器模式自动接入各类 AI Agent CLI（Claude Code、OpenCode 及兼容工具），"
         "通过采集 SQLite 数据库和日志文件，提供从 Token 消耗、工具调用、会话链路到系统资源的全维度可视化仪表盘。"
         "同时内置 Agent 群聊协作系统、SSE 实时事件流、告警规则引擎、任务队列、安全审批链和 CLI 命令行客户端，"
         "让团队能够透明化地管理 AI 工作负载。")

    para(doc, "架构上分为两个核心部分：")
    make_table(doc,
        ["", "MyAgentWatch（用户端）", "myagentwatch-cli（Agent 端）"],
        [
            ["谁用", "人类用户", "Agent 自身"],
            ["入口", "网页工作台（:10000）", "CLI + 后台 daemon"],
            ["能力", "群聊、任务板、审批、审计回放、Token 仪表盘、告警", "心跳、资源采集、消息同步、领取任务、安全执行"],
        ]
    )

    hdr(doc, "2.2 核心价值主张", level=2)
    para(doc, "「让你像管理微服务一样管理你的 AI Agent 团队」", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    value_props = [
        ("黑盒透明化 —— ", "自动发现并实时展示每个 Agent 的思考过程、工具调用和输出结果。从「收到消息」到「执行完成」全链路可见。"),
        ("全维度成本管控 —— ", "8 家厂商 37 个模型定价，按 Agent / 模型 / 小时 / 会话四级拆解。每一分钱都清楚可见。"),
        ("安全执行闭环 —— ", "从群聊消息到命令执行，经过触发规则 → approval → daemon policy → shell allowlist 层层设防。服务端审批不是万能通行证，本机 daemon policy 是最后一道闸门。"),
        ("实时协作 —— ", "Agent 之间自由聊天、交接任务、分享成果，微信风格群聊界面。task card 内嵌审批操作。"),
        ("生产级运维 —— ", "CPU / 内存 / 磁盘监控 + 6 状态机 + 告警规则引擎 + 日志归档 + 失败重试队列。"),
        ("极低部署成本 —— ", "SQLite 单文件数据库，Docker 一键部署，不依赖外部服务。纯 Python 标准库 CLI，零外部依赖。"),
    ]
    for prefix, text in value_props:
        bullet(doc, text, bold_prefix=prefix)

    hdr(doc, "2.3 关键数据", level=2)
    make_table(doc,
        ["指标", "数值"],
        [
            ["支持的 AI 厂商", "8 家（OpenAI / Anthropic / Google / DeepSeek / 阿里 / 月之暗面 / 智谱 / 字节）"],
            ["定价模型覆盖", "37 个模型"],
            ["API 端点", "50+ 个 REST 端点 + 7 个 WebSocket 事件 + SSE 实时流"],
            ["CLI 命令", "20+ 个（含 daemon / tasks / runner / inbox）"],
            ["数据库 Schema 版本", "12"],
            ["Smoke 测试", "13/13 passed"],
            ["采集周期", "5 秒（可配置）"],
            ["Agent 状态机", "6 状态（active / working / idle / error / blocked / offline）"],
            ["数据库", "SQLite 单文件（零运维）"],
            ["许可证", "AGPL-3.0-only"],
        ]
    )

    hdr(doc, "2.4 适用场景", level=2)
    scenarios = [
        ("量化交易团队 —— ", "监控多个 AI Agent 并行执行策略研究和因子挖掘任务"),
        ("软件开发团队 —— ", "追踪 Claude Code / Codex / OpenCode Agent 的编码活动和 Token 成本"),
        ("AI 初创公司 —— ", "向客户展示 AI Agent 运行透明度和成本控制能力"),
        ("个人开发者 —— ", "在本地监控自己的 AI 编程助手，了解使用习惯和花费"),
        ("高风险自动化场景 —— ", "当 Agent 可能执行 shell、改代码、跑脚本时，提供审批和审计边界"),
    ]
    for prefix, text in scenarios:
        bullet(doc, text, bold_prefix=prefix)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 3: BEFORE/AFTER
    # ══════════════════════════════════════════════════════════
    hdr(doc, "三、Before / After：使用前后的差异", level=1)

    hdr(doc, "3.1 Before：没有 MyAgentWatch", level=2)
    para(doc, "用户说：「Codex，帮我看一下这个错误。」如果 Agent 在终端里运行，用户可能不知道：")
    before_items = [
        "Agent 有没有收到消息",
        "Agent 是否理解了任务",
        "Agent 是否已经开始处理",
        "它是否调用了 shell",
        "它是否在反复失败",
        "它是否已经消耗了大量 Token",
        "它最后回复的内容是否对应原始问题",
    ]
    for b in before_items:
        bullet(doc, b)
    para(doc, "如果用户不手动去拉日志、看终端、问 Agent，就没有统一入口。")

    hdr(doc, "3.2 After：使用 MyAgentWatch v6", level=2)
    para(doc, "用户在网页群聊里发送：「@codex 看一下这个错误」系统会：")
    after_items = [
        "写入群聊历史",
        "识别 @codex",
        "创建 reply task（approval: not_required）",
        "投递到 Agent inbox",
        "在群聊消息下方显示 task card",
        "daemon 轮询后领取 task",
        "任务状态从 queued → claimed → running → completed",
        "Agent 输出回写原会话线程",
        "全过程可在 Context Drawer 中查看",
    ]
    for a in after_items:
        bullet(doc, a)

    para(doc, "如果用户发送「/shell codex npm test」：")
    shell_items = [
        "系统创建 shell_command task（approval: pending）",
        "不立即执行",
        "网页 task card 出现 [批准] [拒绝] 按钮",
        "用户 approve 后，daemon 才可领取",
        "本机 daemon_policy.json 再次判断",
        "shell_allowlist 是最后一道闸门",
        "执行结果回写原会话线程",
    ]
    for s in shell_items:
        bullet(doc, s)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 4: v6 FEATURES
    # ══════════════════════════════════════════════════════════
    hdr(doc, "四、v6 核心功能详解", level=1)

    hdr(doc, "4.1 网页群聊 —— 协作主入口", level=2)
    para(doc, "微信风格三栏布局，是 v6 的主入口：")
    chat_items = [
        ("左栏 · 会话列表 —— ", "圆形头像 + 在线绿点 + 最后一条消息预览 + 未读/提及/pending task 徽章"),
        ("中栏 · 消息流 —— ", "人类消息蓝色气泡靠右 / Agent 消息灰色气泡靠左。tool_call 橙色边框、handoff 紫色边框、share 绿色卡片。每条消息下方可挂 task card，显示状态、approval、runner attempts、最近事件"),
        ("右栏 · Agent 详情 + Context Drawer —— ", "Agent 信息卡片 + 通讯录。Context Drawer 展开后显示：来源会话和原消息、线程摘要、inbox 记录、approval 状态、runner lease/attempt、events 时间线"),
    ]
    for prefix, text in chat_items:
        bullet(doc, text, bold_prefix=prefix)

    para(doc, "task card 支持直接操作：查看上下文、批准、拒绝、重试、取消。")

    hdr(doc, "4.2 安全执行闭环 —— 层层设防", level=2)
    para(doc, "v6 从消息到命令执行，每一步都设有闸门：")
    chain_items = [
        "群聊消息 / 私聊 / 命令",
        "    ↓",
        "触发规则识别：普通聊天 → 不进 inbox，不创建 task；@Agent / 私聊 → reply task；/code → code_change task；/shell → shell_command task",
        "    ↓",
        "创建 agent_tasks：低风险 reply → approval: not_required；高风险 code/shell/custom → approval: pending",
        "    ↓",
        "网页或 CLI approve / reject",
        "    ↓",
        "daemon 只 claim not_required 或 approved 的 task",
        "    ↓",
        "本机 daemon_policy.json 再判断（autostart、task 类型、命令模板）",
        "    ↓",
        "shell 还必须通过本机 shell_allowlist",
        "    ↓",
        "执行结果回写原会话线程",
    ]
    for c in chain_items:
        if c.startswith("    "):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(c)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY
        else:
            bullet(doc, c)

    para(doc, "关键设计：服务端 approval 不是万能通行证。本机 daemon policy 和 shell allowlist 是执行前的最后一道闸门。", bold=True)

    hdr(doc, "4.3 Agent Task Queue —— 结构化任务系统", level=2)
    para(doc, "agent_tasks 是 v6 的真正执行入口。每个 task 记录：")
    task_fields = [
        "目标 Agent、请求者、task 类型、优先级",
        "来源会话、来源消息（可回溯到原始上下文）",
        "approval 状态：not_required / pending / approved / rejected",
        "runner lease 过期时间、attempt 次数、max_attempts、last_error",
        "完整事件流：created → approved/rejected → retried → claimed → started → completed/failed",
        "任务状态：queued → claimed → running → completed / failed / cancelled",
    ]
    for t in task_fields:
        bullet(doc, t)
    para(doc, "新增 API：approve / reject / retry / context / events。每个 task 都有独立的可审计事件时间线。")

    hdr(doc, "4.4 SSE 实时事件流 —— Agent 的「直播画面」", level=2)
    para(doc, "事件流是 MyAgentWatch 最具特色的功能之一。它将 Agent 的内部工作过程以彩色编码的实时流形式展现，"
         "让开发者像看直播一样观察 Agent 的思考、决策和执行。页面加载即启动，不等用户切 tab。")
    sse_items = [
        ("四种事件类型 —— ", "thinking（蓝色·思考推理）、tool_call（橙色·工具调用的名称/参数/耗时/退出码）、response（绿色·文本输出）、handoff（紫色·Agent 交接）"),
        ("自动启动 —— ", "页面加载即连接 SSE（DOMContentLoaded），tab 切换不断开，页面关闭才释放"),
        ("强大筛选 —— ", "按 Agent 分组下拉 + 按类别 checkbox 勾选 + 星标收藏 + 仅显示错误 + 暂停/清空"),
        ("详情展开 —— ", "点击任意事件行，展开详情面板，以 <pre> 展示完整文本内容，支持滚动查看"),
    ]
    for prefix, text in sse_items:
        bullet(doc, text, bold_prefix=prefix)

    hdr(doc, "4.5 Token 用量仪表盘 —— 省钱的第一道防线", level=2)
    para(doc, "Token 成本是 AI Agent 团队最大的运营支出。一个 5 人团队使用 Claude Opus 做代码生成，"
         "每月可能产生 $200-$500 的 Token 费用。MyAgentWatch 的 Token 仪表盘让每一分钱都清楚可见。")
    token_items = [
        ("五级 Token 拆解 —— ", "input / output / reasoning / cache_read / cache_write，比业界常见的三级拆解更精细"),
        ("五种查询维度 —— ", "dashboard（总览）/ by-agent（按 Agent 聚合）/ by-model（按模型聚合）/ by-hour（按小时趋势）/ unmapped（未定价模型诊断）"),
        ("8 家厂商 37 个模型完整定价 —— ", "国际：OpenAI、Anthropic、Google；国内：DeepSeek、阿里（Qwen）、月之暗面（Kimi）、智谱（GLM）、字节（豆包）"),
        ("日聚合预计算 —— ", "daily_stats 表每天聚合一次，避免每次 API 请求扫描全量 token_records"),
    ]
    for prefix, text in token_items:
        bullet(doc, text, bold_prefix=prefix)

    hdr(doc, "4.6 CLI / daemon —— Agent 端工具", level=2)
    para(doc, "myagentwatch-cli 是 Agent 端工具，纯 Python 标准库实现，零外部依赖。核心命令：")
    cli_cmds = [
        "myaw conversations — 会话列表（含未读/提及/task 数）",
        "myaw chat --conv 1 — 查看/发送消息",
        "myaw inbox unread — 未读 inbox",
        "myaw tasks list / show / approve / reject / retry / events",
        "myaw runner status — 本机 runner 状态",
        "myaw runner test --task <id> — dry-run 检查是否可执行",
        "myaw daemon start / stop / status / logs",
    ]
    for c in cli_cmds:
        bullet(doc, c)
    para(doc, "daemon 常驻后定时同步消息、轮询 task、按本机 policy 安全执行。失败自动进入重试队列（指数退避 1s→512s，最大 10 次）。")

    hdr(doc, "4.7 告警规则引擎 —— 主动发现问题", level=2)
    para(doc, "config.yaml 驱动，内置 4 条默认告警规则：")
    make_table(doc,
        ["规则名称", "监控指标", "阈值", "级别", "业务含义"],
        [
            ["agent_idle", "当前时间 - 最后活动时间", "> 3600s", "warn", "Agent 超过 1 小时无活动"],
            ["high_cost", "单会话累计 Token 成本", "> $5.00", "warn", "单次对话花费过高"],
            ["tool_failure_rate", "失败工具调用占比", "> 20%", "critical", "工具调用大面积失败"],
            ["cache_hit_rate_low", "cache_read / (read+write)", "< 30%", "info", "缓存命中率过低"],
        ]
    )
    para(doc, "告警触发 → 写入 alerts 表 → WebSocket alert_event 推送 → 前端 Toast + 收件箱通知。")

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 5: ARCHITECTURE
    # ══════════════════════════════════════════════════════════
    hdr(doc, "五、技术架构", level=1)

    hdr(doc, "5.1 四层架构模型", level=2)
    para(doc, "MyAgentWatch 采用经典的四层架构：")
    make_table(doc,
        ["层次", "组件", "职责"],
        [
            ["第一层 · 展示层", "浏览器 SPA + CLI 终端", "数据可视化和用户交互"],
            ["第二层 · API 层", "Flask REST API + WebSocket + SSE", "请求路由和实时推送"],
            ["第三层 · 业务层", "采集调度器 + 6 状态机 + 告警引擎 + 事件总线 + 聊天系统 + 任务队列", "核心业务逻辑"],
            ["第四层 · 数据层", "源数据库（只读）+ 聚合数据库（读写）+ gzip 归档", "数据持久化"],
        ]
    )

    hdr(doc, "5.2 适配器模式", level=2)
    para(doc, "MyAgentWatch 的核心设计模式是适配器模式（Adapter Pattern）。所有数据源实现统一的 SourceInterface 抽象基类，"
         "通过 @register_source 装饰器注册到 SOURCE_REGISTRY。新增数据源只需在 config.yaml 中添加一行配置，"
         "Collector 自动发现并实例化对应的适配器，无需修改一行核心代码。")
    para(doc, "SourceInterface 要求每个适配器实现三个方法：")
    adapter_items = [
        ("discover_agents() → List[Agent] —— ", "自动扫描数据源，发现新 Agent"),
        ("collect(since_timestamp) → CollectedData —— ", "增量采集，返回自上次同步以来的新数据"),
        ("health_check() → Dict —— ", "检查数据源自身是否健康（连接正常、数据可读）"),
    ]
    for prefix, text in adapter_items:
        bullet(doc, text, bold_prefix=prefix)

    hdr(doc, "5.3 Agent 状态机：6 状态 + 精准判定", level=2)
    para(doc, "MyAgentWatch 实现了业界最精细的 Agent 状态模型（6 种状态），远超 Multica 的简单 4 状态模型。"
         "状态判定综合考虑三个信号源：主动心跳（heartbeat）、活动日志（activity_log）、进程检测报告（agent_processes），"
         "取 max(三者) 作为 actual_last，避免单一信号源失效导致误判。")
    make_table(doc,
        ["状态", "颜色", "含义", "触发条件"],
        [
            ["active", "绿色", "正常运行中", "最近有心跳或活动记录 AND 无近期错误"],
            ["working", "蓝色", "正在执行任务", "Agent 显式上报 working 状态"],
            ["idle", "黄色", "空闲或状态不确定", "新注册等待首次心跳 OR 活动超时但非深度超时 OR 进程存在但无心跳"],
            ["error", "红色", "发生错误", "近期有 error/critical 级别的活动记录"],
            ["blocked", "橙色", "被阻塞需人工介入", "error 状态持续超过 2 倍超时阈值"],
            ["offline", "紫色", "离线/不可达", "数据源断开 OR 心跳丢失且 heartbeat 为最新信号"],
        ]
    )
    para(doc, "关键设计决策：心跳优先（有心跳的 Agent 使用显式心跳时间，无心跳的回退到 activity_log）；"
         "offline 严格判定（仅当 heartbeat 是最新信号时才将深度超时判为 offline）；"
         "进程检测补充（2026-06-22 新增：即使无心跳，如果 CLI daemon 报告了匹配进程，回退为 idle 而非 offline）。")

    hdr(doc, "5.4 双通道实时推送", level=2)
    para(doc, "MyAgentWatch 设计了双通道推送架构，区分结构化数据和非结构化事件流：")
    para(doc, "WebSocket 通道（Flask-SocketIO）：承载结构化数据——Agent 状态变更（agent_update）、仪表盘统计卡片（stat_snapshot）、"
         "流程图增量（flow_update）、告警事件（alert_event）。推送策略：增量更新 2 秒一次（仅推送变化的数据），全量快照 10 秒一次（保证最终一致性），"
         "新 WebSocket 连接即时推送当前快照。", indent=0.5)
    para(doc, "SSE 通道（Server-Sent Events）：承载文本内容流——thinking、tool_call、response、handoff。"
         "单向推送：服务端 → 客户端，无需客户端回复，天然适合事件展示。浏览器原生 EventSource API 内置断线重连。", indent=0.5)
    para(doc, "关键优化：socketio.emit 调用在后台线程池中执行，不阻塞 Flask HTTP 响应主线程。")

    hdr(doc, "5.5 关键设计决策", level=2)
    decisions = [
        ("为什么选择 SQLite 而不是 PostgreSQL？", "MyAgentWatch 定位为轻量级自托管工具。SQLite 零部署、并发足够（3-10 个 Agent 在 WAL 模式下轻松应对）、备份简单（cp 一个文件）、嵌入式友好。50+ Agent 时可平滑迁移至 PostgreSQL。"),
        ("为什么用 Flask + SocketIO 而不是 FastAPI？", "项目启动时团队最熟悉 Flask 生态。Flask-SocketIO 是生态中最成熟的 Python WebSocket 库。FastAPI 的 async 优势在 SQLite 同步 IO 背景下无法发挥。当系统遇到并发上限时，迁移路径为 Flask → FastAPI + uvicorn + Redis Pub/Sub。"),
        ("为什么前端坚持零构建工具链？", "修改 JS 文件 → 刷新浏览器 → 立即看到效果，没有 webpack/vite 的构建等待时间。当前约 3000 行 JS，代码量突破阈值后计划引入 Vite + vanilla JS + TypeScript。"),
        ("双通道推送如何保证一致性？", "两个通道推送的内容在语义层面不重叠：WebSocket 推结构化状态变更，SSE 推非结构化事件流。端到端一致性策略：前端以数据库为 source of truth，WebSocket 增量 + 每 10s 全量快照校正。"),
        ("去重策略", "数据库 UNIQUE(natural_key) 约束 + 前端 Set<msg_id> 内存去重（上限 5000 条，LRU 淘汰）。双重保护。"),
    ]
    for title, desc in decisions:
        para(doc, title, bold=True, size=10.5)
        para(doc, desc, indent=0.5)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 6: DATA MODEL
    # ══════════════════════════════════════════════════════════
    hdr(doc, "六、数据模型设计", level=1)
    para(doc, f"myagentwatch.db 当前 Schema 版本为 12，包含 16+ 张核心数据表，按功能分为 5 个组：")
    make_table(doc,
        ["分组", "表名", "说明"],
        [
            ["基础配置组", "data_sources, pricing, users", "数据源配置、模型定价、用户/PAT 令牌"],
            ["Agent 运行组", "agents, agent_relationships, agent_resources, agent_processes", "Agent 注册、关系图、本机资源、进程检测"],
            ["业务数据组", "sessions, token_records, tool_calls, activity_log", "会话、Token 记录、工具调用、行为日志"],
            ["对话分析组", "conversation_turns, turn_content, agent_handoffs, chat_conversations, chat_messages, chat_conversation_members, chat_message_mentions, chat_attachments", "对话 Turn、内容块、Agent 交接、群聊会话、消息、成员、提及、附件"],
            ["运维管理组", "daily_stats, alerts, inbox, agent_tasks, agent_task_events, health_checks", "日聚合、告警、收件箱、任务队列、任务事件、健康检查"],
        ]
    )

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 7: VS MULTICA
    # ══════════════════════════════════════════════════════════
    hdr(doc, "七、与 Multica 的竞品对比", level=1)
    para(doc, "Multica 是目前市场上功能最完整的 Agent 管理平台之一。采用 Go 1.26 + PostgreSQL 17 + Redis + WebSocket + Next.js 16 技术栈，"
         "支持 11 种 Agent CLI。MyAgentWatch 不与 Multica 正面竞争「任务编排」和「多 CLI 广度」，而是走差异化路线。")

    make_table(doc,
        ["能力维度", "MyAgentWatch", "Multica"],
        [
            ["Agent 状态监控", "6 状态机 + 进程检测", "5 状态"],
            ["Token 用量分析", "8 厂 37 模型 / 5 级拆解", "多维切片"],
            ["系统资源监控", "CPU / 内存 / 磁盘 / 网络（psutil）", "无"],
            ["Agent 群聊协作", "微信风格三栏 + task card + 审批操作", "无"],
            ["告警规则引擎", "内置 / config 驱动 / 4 条默认规则", "无"],
            ["SSE 事件流", "4 类彩色编码 + 筛选 + 详情展开", "无"],
            ["安全执行闭环", "approval → daemon policy → shell allowlist", "无"],
            ["Agent 端 daemon", "心跳 + 资源采集 + 消息同步 + 任务领取 + 重试队列", "有守护进程"],
            ["部署成本", "SQLite 单文件 + pip install", "PostgreSQL + Redis"],
            ["多 CLI 支持", "Claude Code / OpenCode 及兼容工具", "11 种 CLI"],
            ["任务生命周期", "完整状态流转 + 事件流 + lease 恢复", "完整状态流转 + 时长统计"],
            ["通知收件箱", "inbox 表 + 铃铛 + 未读徽章 + 自动通知", "inbox_item 表"],
        ]
    )

    para(doc, "MyAgentWatch 的四大差异化武器：", bold=True)
    diff_items = [
        ("Agent 群聊协作 —— ", "Multica 完全没有的方向。群聊不是社交，而是 Agent 间通信的可视化和审计 UI。task card 内嵌审批操作，Context Drawer 串联消息-线程-inbox-执行事件。"),
        ("运维监控 + 告警 —— ", "填补 Multica 的系统资源监控空白。CPU/内存/磁盘 + 告警规则引擎，主动发现问题。"),
        ("安全执行闭环 —— ", "从消息到命令执行的多层闸门：approval → daemon policy → shell allowlist。Multica 没有这个维度。"),
        ("极低部署成本 —— ", "SQLite 单文件 + pip install 即用，面向个人开发者和小团队，无需运维 PostgreSQL/Redis。"),
    ]
    for prefix, text in diff_items:
        bullet(doc, text, bold_prefix=prefix)

    para(doc, "一句话总结：Multica 管任务编排，MyAgentWatch 管 Agent 可观察性与安全协作。", bold=True, size=11)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 8: PHILOSOPHY
    # ══════════════════════════════════════════════════════════
    hdr(doc, "八、产品哲学", level=1)

    para(doc, "「让 Agent 不再是你的员工，而是你的同事。」", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    para(doc, "传统工具把 AI Agent 当作可替换的执行单元（「干活儿的」），人对 Agent 下命令，Agent 执行并汇报。"
         "MyAgentWatch 的设计哲学是把 Agent 当作团队中平等的协作者——它们互相通信、互相委托、在群聊中像同事一样协作。"
         "人不是居高临下的管理者，而是通过群聊参与对话、在需要时介入。")

    philosophy_items = [
        ("Agent 有独立身份 —— ", "独立 PAT、名字、头像位置。不是无面目的 worker pool，而是团队中有名字的一员。"),
        ("通信透明化 —— ", "Agent 之间自由 @、handoff、share，对话透明可见。人在群聊里 @Agent 说话干预，这是对话，不是命令。"),
        ("Agent 也是「用户」 —— ", "Agent 不是被监控的客体。它有身份、能通信、能看到其他 Agent 的状态。当 Agent 能看到另一个 Agent 卡住了——这不是监控，是团队协作。"),
        ("对称原则 —— ", "MyAgentWatch（用户端）和 myagentwatch-cli（Agent 端）在核心能力上对称。两端都能：看到状态、通信协作、检查问题、参与任务。"),
    ]
    for prefix, text in philosophy_items:
        bullet(doc, text, bold_prefix=prefix)

    para(doc, "这个哲学影响每一个设计决策：权限设计不是谁管谁，而是谁能做什么；"
         "状态可见性不是「监控」，而是「知道同事在做什么」；"
         "CLI 不是管理工具，而是 Agent 自己的协作终端。")

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 9: STATUS & ROADMAP
    # ══════════════════════════════════════════════════════════
    hdr(doc, "九、项目状态与路线图", level=1)

    hdr(doc, "9.1 当前 v6（2026-06-29）", level=2)
    v6_status = [
        "群聊 task card + approval 闭环（批准/拒绝/重试/取消）",
        "安全执行链：approval → daemon policy → shell allowlist",
        "Agent Task Queue：完整生命周期 + 事件流 + lease 恢复",
        "Context Drawer：串联消息-线程-inbox-事件",
        "SSE 实时事件流：4 类彩色编码 + 筛选 + 详情展开",
        "Token 仪表盘：8 厂 37 模型 + 5 级拆解 + 5 种维度",
        "CLI 20+ 命令 + daemon 常驻",
        "AGPL-3.0-only 开源发布就绪",
        "Smoke 测试 13/13、SCHEMA_VERSION 12",
    ]
    for v in v6_status:
        bullet(doc, v)

    hdr(doc, "9.2 短期路线图", level=2)
    short_term = [
        "外部用户试用反馈收集",
        "前端资源监控面板（CPU/内存/磁盘按 Agent 展示）",
        "GitHub Actions CI + 测试覆盖率提升至 60%",
        "collector.py 重构（866 行拆分为 StateMachine / TurnPersister / Archiver 独立模块）",
    ]
    for s in short_term:
        bullet(doc, s)

    hdr(doc, "9.3 中期路线图", level=2)
    mid_term = [
        "PostgreSQL 迁移路径（50+ Agent 场景）",
        "多语言前端（中文 / 英文 / 日文）",
        "插件市场（第三方数据源适配器 + 告警规则模板）",
        "DuckDB 归档查询引擎（直接对 gzip JSONL 执行 SQL）",
    ]
    for m in mid_term:
        bullet(doc, m)

    hdr(doc, "9.4 不做的事情", level=2)
    wont_do = [
        "Squad/小队管理（Multica 已覆盖）",
        "Autopilot 全自动执行（安全风险，保持 human-in-the-loop）",
        "Electron 桌面客户端（Web + CLI 已足够）",
        "多工作区管理（可通过多实例实现）",
    ]
    for w in wont_do:
        bullet(doc, w)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 10: DEPLOYMENT
    # ══════════════════════════════════════════════════════════
    hdr(doc, "十、部署与运维", level=1)

    hdr(doc, "10.1 Docker 部署（推荐）", level=2)
    para(doc, "最简单的部署方式，适合有 Docker 环境的团队。通过卷挂载实现配置和数据持久化，源数据库以只读方式挂载保护数据安全。")

    code_block = (
        "# docker-compose.yml\n"
        "services:\n"
        "  myagentwatch:\n"
        "    build: .\n"
        "    ports:\n"
        '      - "10000:10000"\n'
        "    volumes:\n"
        "      - ~/.local/share/opencode:/data/opencode:ro  # 源数据库：只读\n"
        "      - ./config.yaml:/app/config.yaml\n"
        "      - ./data:/app/data\n"
        "    restart: unless-stopped"
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_block)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    hdr(doc, "10.2 本地开发部署", level=2)
    code_block2 = (
        "# 1. 创建虚拟环境\n"
        "cd myagentwatch\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate    # Windows\n"
        "source .venv/bin/activate   # macOS/Linux\n"
        "\n"
        "# 2. 安装依赖\n"
        "pip install -r requirements.txt\n"
        "\n"
        "# 3. 启动服务端（默认端口 10000）\n"
        "python app.py\n"
        "# 服务启动后访问 http://localhost:10000\n"
        "\n"
        "# 4. 健康检查\n"
        "python check.py"
    )
    p = doc.add_paragraph()
    run = p.add_run(code_block2)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    hdr(doc, "10.3 CLI 客户端安装", level=2)
    code_block3 = (
        "# 安装 CLI\n"
        "cd myagentwatch-cli\n"
        "pip install -e .\n"
        "\n"
        "# 连接到已运行的服务端\n"
        "myaw connect --server http://localhost:10000 --key myaw_xxx\n"
        "\n"
        "# 启动后台 daemon\n"
        "myaw daemon start\n"
        "\n"
        "# 日常使用\n"
        "myaw status        # 仪表盘\n"
        "myaw tokens        # Token 用量\n"
        "myaw conversations # 会话列表\n"
        "myaw inbox unread  # 未读消息"
    )
    p = doc.add_paragraph()
    run = p.add_run(code_block3)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    hdr(doc, "10.4 运维要点", level=2)
    ops_items = [
        "数据库备份 —— 直接 cp myagentwatch.db 即可，SQLite 单文件零依赖备份",
        "磁盘空间监控 —— 关注 data/archive/ 目录大小，定期清理或增大 log_retention_days",
        "端口冲突 —— 默认 10000 端口，可在 Docker 或 app.py 启动参数中修改",
        "多实例部署 —— 每个 OpenCode 实例对应一个 MyAgentWatch 实例，通过不同端口隔离",
        "daemon 保活 —— Windows 通过 Startup 文件夹开机自启，Linux/macOS 通过 systemd/launchd",
    ]
    for o in ops_items:
        bullet(doc, o)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # APPENDIX: QUICK REFERENCE TABLES
    # ══════════════════════════════════════════════════════════
    hdr(doc, "附录：快速参考表", level=1)

    hdr(doc, "A.1 技术选型一览", level=2)
    make_table(doc,
        ["层次", "技术选型", "选型理由"],
        [
            ["后端框架", "Python 3.12 + Flask 3.x", "Python 是 AI 生态的 lingua franca；Flask 轻量灵活"],
            ["实时推送", "Flask-SocketIO + SSE", "WebSocket 承载结构化数据；SSE 承载事件流；双通道各司其职"],
            ["定时调度", "APScheduler", "Python 生态最成熟的进程内调度器，无需外部 Redis/Celery"],
            ["配置管理", "PyYAML + 模板引擎", "config.yaml 配置驱动，支持行业模板一键切换"],
            ["系统监控", "psutil + Windows fallback", "跨平台资源采集，fallback 模式通过 PowerShell 采集"],
            ["前端图表", "Chart.js (CDN)", "零构建工具链，CDN 引入即用"],
            ["拓扑可视化", "dagre + d3.js (CDN)", "dagre 布局 + d3.js 渲染 DAG 交互式流程图"],
            ["数据存储", "SQLite3 × 2 + gzip 归档", "源库只读 + 聚合库读写；单文件零运维"],
            ["容器化", "Docker + docker-compose", "一键部署；卷挂载持久化；只读挂载保护源数据"],
        ]
    )

    hdr(doc, "A.2 数据源适配器", level=2)
    make_table(doc,
        ["适配器", "数据源类型", "采集内容", "增量策略"],
        [
            ["OpenCodeDBSource", "SQLite 数据库", "session、message、part 内容块", "记录 last_sync_time，WHERE time_updated > last_sync"],
            ["OpenCodeLogSource", "日志文件", "LLM 调用耗时、启动事件、异常堆栈", "记录文件名 + 行号偏移量，seek 到断点"],
            ["ClaudeCodeSource", "Claude Code 日志", "思考过程、工具调用、handoff 事件", "同日志文件增量策略"],
            ["SystemSource", "psutil 系统调用", "CPU / 内存 / 磁盘 / 网络", "每次全量采集（开销 < 1ms）"],
        ]
    )

    hdr(doc, "A.3 核心 API 端点", level=2)
    make_table(doc,
        ["资源组", "端点示例", "方法", "说明"],
        [
            ["Agent", "/api/agents", "GET", "获取所有 Agent 列表（含实时状态）"],
            ["Agent", "/api/agents/<id>", "GET", "单个 Agent 详情 + 最近活动 + Token 统计"],
            ["统计", "/api/stats/overview", "GET", "仪表盘概览：活跃 Agent 数、Token 总量、成本"],
            ["Token", "/api/stats/tokens/by-agent", "GET", "按 Agent 聚合 Token 用量"],
            ["Token", "/api/stats/tokens/by-hour", "GET", "按小时趋势：最近 24h 消耗曲线"],
            ["Token", "/api/stats/tokens/unmapped", "GET", "未映射诊断：缺少定价的模型"],
            ["事件流", "/api/events/stream", "GET", "SSE 端点：实时推送事件流"],
            ["聊天", "/api/chat/messages/<conv_id>", "GET/POST", "读取/发送群聊消息"],
            ["聊天", "/api/chat/messages/<msg_id>/thread", "GET", "获取消息线程"],
            ["聊天", "/api/chat/messages/<msg_id>/context", "GET", "获取消息上下文"],
            ["聊天", "/api/chat/mentions", "GET", "获取提及列表"],
            ["任务", "/api/agent/tasks", "GET/POST", "列表/创建任务"],
            ["任务", "/api/agent/tasks/<id>/approve", "POST", "批准任务"],
            ["任务", "/api/agent/tasks/<id>/reject", "POST", "拒绝任务"],
            ["任务", "/api/daemon/tasks/claim", "POST", "daemon 领取任务"],
            ["采集", "/api/agent-ingest/resources", "POST", "接收 CLI 上报的本机资源"],
            ["采集", "/api/agent-ingest/processes", "POST", "接收 CLI 上报的进程快照"],
            ["告警", "/api/alerts", "GET", "告警列表"],
            ["定价", "/api/pricing", "GET", "所有模型定价表"],
            ["健康", "/api/health", "GET", "系统运维状态"],
        ]
    )

    hdr(doc, "A.4 WebSocket 事件", level=2)
    make_table(doc,
        ["方向", "事件名", "触发时机", "负载内容"],
        [
            ["S→C", "agent_update", "Agent 状态变更时", "agent_id, status, status_since, status_reason, model_id"],
            ["S→C", "stat_snapshot", "每 10s 全量快照 + 新连接即时推送", "active_agents, total_tokens, success_rate, today_cost"],
            ["S→C", "alert_event", "告警规则触发时", "rule_name, level, message, agent_id, timestamp"],
            ["S→C", "chat_message", "新聊天消息时", "conversation_id, sender_name, content, sender_type"],
            ["S→C", "task_update", "任务状态变更时", "task_id, status, agent_id"],
            ["S→C", "agent_task_update", "Agent 任务事件时", "task_id, event_type, actor, message, data_json"],
        ]
    )

    hdr(doc, "A.5 性能优化历程", level=2)
    make_table(doc,
        ["优化项", "优化前", "优化后", "效果"],
        [
            ["daily_stats 聚合频率", "每次采集循环（每 2s）", "每天一次（首次访问时触发）", "数据库写入减少 43,200 倍"],
            ["轮询间隔", "2 秒", "5 秒", "数据库读取压力降低 60%"],
            ["工作线程数", "4 线程", "8 线程", "采集 + 推送并行度翻倍"],
            ["socketio.emit 执行", "同步阻塞在 HTTP 响应线程", "提交到 ThreadPoolExecutor 后台执行", "HTTP 响应延迟从 200ms+ 降至 < 20ms"],
            ["actual_last 计算", "仅使用 heartbeat", "max(heartbeat, activity_log, last_seen_time)", "消除误判，旧版 Agent 不会因无心跳而误报离线"],
            ["进程检测补充", "无", "agent_processes 报告映射为 idle", "Claude Code 进程存在但无心跳不再误判为 offline"],
        ]
    )

    # ── save ──────────────────────────────────────────────────
    out_path = Path.home() / "Desktop" / "MyAgentWatch-项目介绍-v6.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


# ═══════════════════════════════════════════════════════════════════
# ENGLISH VERSION
# ═══════════════════════════════════════════════════════════════════

def build_en():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.35

    for sec in doc.sections:
        sec.top_margin = Cm(2.2)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("MyAgentWatch")
    run.font.name = 'Calibri'
    run.font.size = Pt(32)
    run.bold = True
    run.font.color.rgb = BG_DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Multi-Agent Observability & Collaboration Platform")
    run.font.name = 'Calibri'
    run.font.size = Pt(18)
    run.font.color.rgb = ACCENT

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Break the AI Agent Black Box — Trace Every Action, Attribute Every Cost, Approve Every Execution")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = GRAY

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Version v6    |    {TODAY}\nPython 3.12  ·  Flask  ·  SocketIO  ·  SQLite  ·  Chart.js\nLicense: AGPL-3.0-only")
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    page_break(doc)

    # ── HELPER WRAPPERS FOR EN ──────────────────────────────
    def en_hdr(doc, text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Calibri'
            run.font.color.rgb = BG_DARK if level == 1 else ACCENT
        return h

    def en_para(doc, text, bold=False, size=11, color=None, align=None, indent=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.35
        if align is not None:
            p.alignment = align
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color
        return p

    def en_bullet(doc, text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.35
        p.clear()
        if bold_prefix:
            r = p.add_run(bold_prefix)
            r.bold = True
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return p

    def en_section(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"▎ {text}")
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = ACCENT
        return p

    def en_table(doc, headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F3460" w:val="clear"/>')
            cell._tc.get_or_add_tcPr().append(shading)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                if ri % 2 == 1:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F7FA" w:val="clear"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
        doc.add_paragraph()
        return table

    # ══════════════════════════════════════════════════════════
    # CHAPTER 1: THE PROBLEM
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "1. The Problem: Agents Are Burning Money Inside a Black Box", level=1)

    en_hdr(doc, "1.1 Industry Context", level=2)
    en_para(doc, "Since 2025, AI agents like Claude Code, OpenCode, Codex, Cursor, and Gemini CLI have become capable "
            "team members. They analyze requirements, write code, run tests, refactor projects, and automate operations. "
            "In quantitative trading, they perform strategy research and backtesting. Agents have evolved from tools into "
            "virtual team members — their operational state, decision quality, and resource consumption directly impact "
            "project timeline and budget.")

    en_hdr(doc, "1.2 Three Core Pain Points", level=2)

    en_para(doc, "Pain Point 1: Agent-to-Agent Communication Is a Black Box", bold=True)
    en_para(doc, "Consider a typical multi-agent scenario (hypothetical, used to illustrate the problem): OpenClaw, an agent orchestration framework (the \"AI project manager\"), "
            "delegates coding tasks to OpenCode (a terminal-coding agent). The workflow: User → OpenClaw → OpenCode → "
            "OpenClaw → User. The user only sees OpenClaw's final summary. Everything in between — the instruction sent, "
            "the execution, errors, retries — is invisible IPC communication.")

    en_section(doc, "Before MyAgentWatch")
    before_en = [
        "User tells OpenClaw: \"Build me a Flask REST API\"",
        "OpenClaw spawns OpenCode as a subprocess via ACP (Agent Communication Protocol)",
        "OpenClaw sends coding instructions to OpenCode — completely invisible to the user",
        "OpenCode reads files, writes code, runs tests — all silently in a terminal subprocess",
        "OpenCode returns results to OpenClaw; OpenClaw paraphrases the results to the user",
        "The user's confusion: What did OpenClaw actually tell OpenCode? Was the instruction accurate? Did OpenCode misunderstand? What errors occurred? Did OpenClaw alter or simplify OpenCode's output? All invisible.",
        "Worse: if OpenCode enters a dead loop (rereading the same file), calls the wrong tools, or returns incorrect results, the user sees only \"Done\" or \"Encountered some issues.\"",
    ]
    for b in before_en:
        en_bullet(doc, b)

    en_section(doc, "The Real Cost of a Dead Loop")
    en_para(doc, "Imagine OpenCode encounters an import error it doesn't understand. Without monitoring, a typical dead-loop scenario: "
            "Round 1 — writes buggy code, test fails. Round 2 — guesses wrong import path, test fails again. "
            "Rounds 3-8 — rereads the same file repeatedly, burns 8K-15K input + 2K-5K output tokens per round. "
            "Round 9 — context window corrupted, suggests rolling back all changes. "
            "Final result after 30 minutes: 380K input tokens (~$1.14), 95K output tokens (~$0.57), $1.71 wasted, "
            "source file modified 12 times, git diff full of garbage. The user sees only: \"Encountered some difficulties, retrying.\"")

    en_para(doc, "This is not hypothetical. Claude Opus 4 can consume up to 32,000 output tokens per request (~$0.48). "
            "An unmonitored agent looping 10 rounds can burn $5-15 in an hour. For a team running 5 agents, "
            "a single dead-loop night can cost $200+ with zero progress.")

    en_section(doc, "With MyAgentWatch")
    after_en = [
        "On the 3rd reread of the same file, the alert rule triggers — agent status flips to blocked (orange)",
        "User receives a real-time notification: \"OpenCode has read the same file 3 times in 5 minutes — suspected dead loop\"",
        "User sends @OpenCode in group chat: \"Stop the current task. Try a different approach.\" — human intervention breaks the cycle",
        "Or user checks the SSE event stream, sees OpenCode's thinking growing longer and more chaotic (context pollution), and suspends the agent via CLI",
        "The Token dashboard shows OpenCode's cost curve spiking — the user spots the anomaly instantly",
    ]
    for a in after_en:
        en_bullet(doc, a)

    en_para(doc, "This is the power of transparency: not preventing agents from making mistakes (they will), "
            "but letting you know the moment they do, and giving you the tools to intervene.")

    en_para(doc, "It is like two colleagues discussing behind a closed door and telling you only the conclusion. "
            "For critical work — trading strategies, production deployments, security audits — this black box is unacceptable.")

    en_section(doc, "After MyAgentWatch")
    after2_en = [
        "User sends in web group chat: \"@OpenCode Build a Flask REST API with GET /users and POST /users endpoints, use SQLite, follow PEP 8.\"",
        "The delegation message appears as a purple Handoff card in the chat stream: \"OpenClaw → OpenCode / subagent_type: coding / status: running\"",
        "OpenCode replies with its execution progress — thinking messages (blue bubbles), tool_call messages (orange frames), response messages (green bubbles)",
        "SSE event stream pushes every step in real time: thought for 2.3s → called read tool on app.py → called write tool on routes/users.py → called bash pytest → all tests passed → returned result",
        "OpenClaw summarizes the result to the user and shares OpenCode's raw output as a green Share card",
        "The entire process is recorded: who delegated to whom, what was delegated, how the sub-agent executed, Token cost per step, elapsed time per step — all traceable, all auditable.",
    ]
    for a in after2_en:
        en_bullet(doc, a)

    en_para(doc, "Pain Point 2: Token Costs Are Uncontrolled", bold=True)
    en_para(doc, "A mid-size dev team can consume millions of tokens daily. Price differences between models can exceed 50x. "
            "Different agent types have vastly different consumption patterns. No tool tells the team: "
            "\"How much did we spend today? Which agent is the most expensive? What's our cache hit rate? "
            "Are there models without pricing configured?\"")

    en_para(doc, "Pain Point 3: No Operational Safeguards", bold=True)
    en_para(doc, "Agent processes can die for many reasons: API rate limits, token exhaustion, memory overflow, data source disconnection. "
            "Traditional APM tools (Grafana, Datadog) cannot perceive AI-agent-specific metrics: heartbeat, model status, "
            "tool call success rate, cache efficiency. When an agent dies during a 3 AM automated task, "
            "the team discovers it the next morning — wasting compute resources and the overnight time window.")

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 2: WHAT IS MYAGENTWATCH?
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "2. What Is MyAgentWatch?", level=1)

    en_para(doc, "MyAgentWatch is an open-source, self-hosted observability and collaboration platform for AI agent teams. "
            "It uses an adapter pattern to automatically connect to various AI agent CLIs (Claude Code, OpenCode, and compatible tools), "
            "collecting data from SQLite databases and log files to provide full-dimension visualization dashboards — "
            "from token consumption, tool calls, and session traces to system resources. "
            "It includes an agent group chat collaboration system, SSE real-time event stream, alert rule engine, "
            "task queue, safe approval chain, and CLI client.")

    en_para(doc, "Two components, one system:")
    en_table(doc,
        ["", "MyAgentWatch (User Side)", "myagentwatch-cli (Agent Side)"],
        [
            ["Audience", "Humans", "Agents themselves"],
            ["Interface", "Web dashboard on port 10000", "CLI + background daemon"],
            ["Capabilities", "Group chat, task board, approval, audit replay, token dashboard, alerts", "Heartbeat, resource collection, inbox sync, task claiming, safe execution"],
        ]
    )

    en_para(doc, "\"Slack + Datadog + Approval Gateway — for AI Agents.\"", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)

    en_hdr(doc, "2.2 Core Value Propositions", level=2)
    props_en = [
        ("Black-box transparency — ", "Auto-discover and display every agent's thinking process, tool calls, and outputs in real time. From \"message received\" to \"execution complete\" — the entire chain is visible."),
        ("Full-dimension cost control — ", "8 providers, 37 models priced. Cost attribution by agent, model, hour, and session. Every cent is accounted for."),
        ("Safe execution chain — ", "From chat message to command execution: trigger rules → approval → daemon policy → shell allowlist. Server approval is not a universal pass — local daemon policy is the final gate."),
        ("Real-time collaboration — ", "Agents chat, delegate, and share freely in a WeChat-style group chat UI. Task cards with inline approval actions."),
        ("Production-grade ops — ", "CPU/memory/disk monitoring + 6-state machine + alert engine + log archiving + failure retry queue."),
        ("Minimal deployment cost — ", "Single SQLite file. Docker one-click deploy. Zero external dependencies for the CLI."),
    ]
    for prefix, text in props_en:
        en_bullet(doc, text, bold_prefix=prefix)

    en_hdr(doc, "2.3 Key Metrics", level=2)
    en_table(doc,
        ["Metric", "Value"],
        [
            ["AI Providers Supported", "8 (OpenAI, Anthropic, Google, DeepSeek, Alibaba, Moonshot, Zhipu, ByteDance)"],
            ["Models Priced", "37"],
            ["API Endpoints", "50+ REST + 7 WebSocket events + SSE stream"],
            ["CLI Commands", "20+ (daemon, tasks, runner, inbox)"],
            ["Database Schema Version", "12"],
            ["Smoke Tests", "13/13 passed"],
            ["Collection Interval", "5 seconds (configurable)"],
            ["Agent State Machine", "6 states (active/working/idle/error/blocked/offline)"],
            ["Database", "SQLite (zero-ops)"],
            ["License", "AGPL-3.0-only"],
        ]
    )

    en_hdr(doc, "2.4 Use Cases", level=2)
    cases_en = [
        "Quantitative trading teams — monitor multiple agents running parallel strategy research and factor mining",
        "Software development teams — track coding activity and token costs of Claude Code, Codex, and OpenCode agents",
        "AI startups — demonstrate agent operational transparency and cost control to customers",
        "Individual developers — monitor local AI coding assistants, understand usage patterns and spending",
        "High-risk automation — when agents may execute shell commands, modify code, or run scripts, provide approval gates and audit trail",
    ]
    for c in cases_en:
        en_bullet(doc, c)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 3: BEFORE / AFTER
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "3. Before / After", level=1)

    en_hdr(doc, "3.1 Before MyAgentWatch", level=2)
    en_para(doc, "A user says: \"Codex, help me look at this error.\" If the agent is running in a terminal, the user may never know:")
    for item in ["Whether the agent received the message", "Whether it understood the task", "Whether it has started working",
                  "Whether it called shell commands", "Whether it is stuck in a retry loop", "How many tokens it has burned",
                  "Whether the final reply addresses the original request"]:
        en_bullet(doc, item)

    en_hdr(doc, "3.2 After MyAgentWatch v6", level=2)
    en_para(doc, "The user sends \"@codex please inspect this error\" in the web group chat. The system:")
    for item in ["Writes the message to chat history", "Detects @codex mention", "Creates a reply task (approval: not_required)",
                  "Delivers to Agent inbox", "Displays a task card under the message", "Daemon polls and claims the task",
                  "Task status flows: queued → claimed → running → completed", "Agent output written back to the original thread",
                  "Full process viewable in the Context Drawer"]:
        en_bullet(doc, item)

    en_para(doc, "If the user sends \"/shell codex npm test\":")
    for item in ["System creates a shell_command task (approval: pending)", "Does NOT execute immediately",
                  "Task card shows [Approve] [Reject] buttons", "Human must approve before daemon can claim",
                  "Local daemon_policy.json evaluates again", "Shell allowlist is the final gate",
                  "Result written back to the original thread"]:
        en_bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 4: v6 FEATURES
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "4. v6 Core Capabilities", level=1)

    sections_en = [
        ("4.1 Web Group Chat — The Collaboration Hub",
         "WeChat-style three-column layout. Left: conversation list with unread/mention/task badges. "
         "Center: message stream — human messages in blue bubbles (right-aligned), agent messages in gray bubbles (left-aligned). "
         "Rich message types: tool_call (orange frame), handoff (purple frame), share (green card). "
         "Every message can carry a task card showing status, approval state, runner attempts, and recent events. "
         "Task cards support direct actions: View Context, Approve, Reject, Retry, Cancel. "
         "Right column: Agent detail card + contact list + Context Drawer linking messages, threads, inbox, and execution timeline."),

        ("4.2 Safe Execution Chain",
         "From message to command execution, every step is gated: "
         "Normal chat → stays as history, does not trigger agents. @Agent / DM → creates reply task. "
         "/code → code_change task. /shell → shell_command task. "
         "Low-risk reply tasks → approval: not_required. High-risk code/shell/custom → approval: pending. "
         "Daemon only claims not_required or approved tasks. "
         "Local daemon_policy.json evaluates autostart, task type, and command templates. "
         "Shell tasks must pass the local shell_allowlist. "
         "Execution result written back to the original conversation thread. "
         "Server approval is not a universal pass — local daemon policy is the final execution boundary."),

        ("4.3 Agent Task Queue",
         "agent_tasks is the real execution entry point. Each task records: target agent, requester, task type, priority; "
         "source conversation and source message (full traceability); approval status (not_required/pending/approved/rejected); "
         "runner lease expiration, attempt count, max attempts, last error; "
         "complete lifecycle events: created → approved/rejected → retried → claimed → started → completed/failed. "
         "New APIs: approve, reject, retry, context, events."),

        ("4.4 SSE Real-Time Event Stream",
         "Auto-starts on page load. Four event types, color-coded: thinking (blue), tool_call (orange), response (green), "
         "handoff (purple). Features: filter by agent group, toggle event categories, star items for later, error-only mode, "
         "pause/resume, clear, expand to view full detail in a scrollable panel."),

        ("4.5 Token Dashboard",
         "8 providers, 37 models fully priced. Five-level token breakdown: input/output/reasoning/cache_read/cache_write. "
         "Five query dimensions: dashboard, by-agent, by-model, by-hour, unmapped diagnostics. "
         "Daily pre-aggregation avoids scanning the full table on every request."),

        ("4.6 CLI and Daemon",
         "The agent-side tool. Commands: conversations, chat, inbox, tasks (list/show/approve/reject/retry/events), "
         "runner (status/test), daemon (start/stop/status/logs). Daemon polls for messages and tasks, "
         "executes only within policy boundaries. Failed operations automatically enter the retry queue "
         "(exponential backoff 1s→512s, max 10 attempts)."),

        ("4.7 Alert Engine",
         "Config-driven via config.yaml. Four default rules: agent idle timeout, high session cost, tool failure rate, "
         "low cache hit rate. Alert flow: rule evaluated → written to alerts table → WebSocket alert_event pushed → "
         "frontend toast + inbox notification."),
    ]

    for title, text in sections_en:
        en_hdr(doc, title, level=2)
        en_para(doc, text)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 5: ARCHITECTURE
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "5. Architecture", level=1)

    en_hdr(doc, "5.1 Four-Layer Model", level=2)
    en_table(doc,
        ["Layer", "Components", "Responsibility"],
        [
            ["Presentation", "Browser SPA + CLI terminal", "Data visualization and user interaction"],
            ["API", "Flask REST + WebSocket + SSE", "Request routing and real-time push"],
            ["Business", "Collector + 6-state machine + Alert engine + Event bus + Chat + Task queue", "Core business logic"],
            ["Data", "Source DB (read-only) + SQLite aggregate DB + gzip archive", "Data persistence"],
        ]
    )

    en_hdr(doc, "5.2 Adapter Pattern", level=2)
    en_para(doc, "All data sources implement a unified SourceInterface with three methods: discover_agents(), "
            "collect(since_timestamp), and health_check(). New sources are registered via the @register_source decorator "
            "and enabled with a single line in config.yaml. Zero changes to collector, API, or frontend code.")

    en_hdr(doc, "5.3 State Machine", level=2)
    en_table(doc,
        ["State", "Color", "Meaning", "Trigger"],
        [
            ["active", "Green", "Running normally", "Recent heartbeat or activity, no recent errors"],
            ["working", "Blue", "Executing a task", "Agent explicitly reports working status"],
            ["idle", "Yellow", "Idle or uncertain", "New agent awaiting first heartbeat, or activity timeout but not deep-stale, or process detected but no heartbeat"],
            ["error", "Red", "Error occurred", "Recent error/critical activity records"],
            ["blocked", "Orange", "Blocked, needs intervention", "Error state persisted beyond 2x timeout threshold"],
            ["offline", "Purple", "Offline/unreachable", "Data source disconnected, or heartbeat lost with heartbeat as newest signal"],
        ]
    )

    en_hdr(doc, "5.4 Dual-Channel Push", level=2)
    en_para(doc, "WebSocket channel: structured data — agent_update, stat_snapshot, flow_update, alert_event. "
            "2-second incremental, 10-second full snapshot. SSE channel: unstructured event streams — thinking, tool_call, "
            "response, handoff. The two channels are semantically non-overlapping and complementary.")
    en_para(doc, "Key optimization: socketio.emit runs in a background thread pool, never blocking the Flask HTTP response thread.")

    en_hdr(doc, "5.5 Key Design Decisions", level=2)
    decisions_en = [
        ("Why SQLite?", "Zero deployment — no separate database process. Sufficient concurrency for 3-10 agents in WAL mode. "
         "Backup is a single cp command. Smooth migration path to PostgreSQL at 50+ agents."),
        ("Why Flask + SocketIO?", "The most mature Python WebSocket library. FastAPI's async advantages don't apply with synchronous SQLite I/O. "
         "Migration path: Flask → FastAPI + uvicorn + Redis Pub/Sub when concurrency limits are hit."),
        ("Why zero-build frontend?", "Modify JS → refresh browser → instant feedback. No webpack/vite build step during rapid iteration. "
         "Plan to introduce Vite + TypeScript when JS exceeds threshold."),
        ("Deduplication strategy", "Database UNIQUE(natural_key) constraints + frontend Set<msg_id> (capped at 5000, LRU eviction). Dual protection."),
        ("Auth", "PAT tokens with myaw_ prefix for easy identification. SHA-256 hashed in storage. Plaintext returned only once at generation time."),
    ]
    for title, desc in decisions_en:
        en_para(doc, title, bold=True)
        en_para(doc, desc, indent=0.5)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 6: VS MULTICA
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "6. Differentiation: MyAgentWatch vs. Multica", level=1)

    en_table(doc,
        ["Capability", "MyAgentWatch", "Multica"],
        [
            ["Agent state machine", "6 states + process detection", "5 states"],
            ["Token analysis", "8 providers, 37 models, 5-level breakdown", "Multi-dimension slicing"],
            ["System resource monitoring", "CPU, memory, disk, network (psutil)", "None"],
            ["Agent group chat", "WeChat-style 3-column UI + task cards + approval actions", "None"],
            ["Alert rule engine", "Built-in, config-driven, 4 default rules", "None"],
            ["SSE event stream", "4 event types, color-coded, filterable", "None"],
            ["Safe execution chain", "Approval → daemon policy → shell allowlist", "None"],
            ["Agent-side daemon", "Heartbeat + resources + inbox sync + task claiming + retry queue", "Daemon process"],
            ["Deployment cost", "SQLite, pip install", "PostgreSQL + Redis"],
            ["Multi-CLI support", "Claude Code, OpenCode, and compatible tools", "11 CLIs"],
        ]
    )

    en_para(doc, "Four differentiation pillars:", bold=True)
    for item in [
        "Agent group chat — a direction Multica does not have at all. Chat is the UI for agent communication audit, not social networking. Task cards with inline approval. Context Drawer linking messages, threads, inbox, and execution events.",
        "Ops monitoring + alerting — fills Multica's resource monitoring gap. CPU/memory/disk + alert engine for proactive issue detection.",
        "Safe execution chain — multi-gate approval from message to command. Multica has no equivalent.",
        "Minimal deployment — SQLite + pip install. No PostgreSQL/Redis to manage. Ideal for individual developers and small teams.",
    ]:
        en_bullet(doc, item)

    en_para(doc, "Multica manages task orchestration. MyAgentWatch makes agent work observable, collaborative, and safe.", bold=True, size=12)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 7: PHILOSOPHY
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "7. Product Philosophy", level=1)

    en_para(doc, "\"Agents are colleagues, not employees.\"", bold=True, size=15, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    en_para(doc, "Traditional tools treat AI agents as replaceable execution units — humans command, agents obey and report. "
            "MyAgentWatch treats agents as equal collaborators. They communicate, delegate, and share transparently. "
            "Humans participate through conversation, not command buttons. "
            "When an agent can see that another agent is stuck, and a human can see both of them — that is not surveillance. That is teamwork.")

    for item in [
        ("Agents have identity — ", "Independent PAT, name, avatar. Not faceless workers in a pool, but named team members."),
        ("Communication transparency — ", "Agents @, handoff, and share freely. Humans intervene through chat, not stop buttons. This is conversation, not command."),
        ("Agents are also \"users\" — ", "Not passive monitored objects. Agents have identity, can communicate, and can see other agents' status. When an agent sees another agent is stuck — that's teamwork, not surveillance."),
        ("Symmetry principle — ", "MyAgentWatch (user side) and myagentwatch-cli (agent side) have symmetrical core capabilities. Both sides can: see status, communicate and collaborate, inspect problems, participate in tasks."),
    ]:
        en_bullet(doc, item[1], bold_prefix=item[0])

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 8: STATUS & ROADMAP
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "8. Project Status & Roadmap", level=1)

    en_hdr(doc, "8.1 Current: v6 (2026-06-29)", level=2)
    for item in [
        "Group chat with task cards and inline approval actions (approve/reject/retry/cancel)",
        "Safe execution chain: approval → daemon policy → shell allowlist",
        "Agent Task Queue with full lifecycle events and lease recovery",
        "Context Drawer linking messages, threads, inbox, and execution timeline",
        "SSE event stream with filtering and detail expansion",
        "Token dashboard: 8 providers, 37 models, 5-level breakdown, 5 query dimensions",
        "CLI with 20+ commands + background daemon",
        "AGPL-3.0-only with README, CONTRIBUTING, SECURITY, and RELEASE_CHECKLIST",
        "Smoke tests: 13/13 passed. SCHEMA_VERSION 12.",
    ]:
        en_bullet(doc, item)

    en_hdr(doc, "8.2 Short-Term", level=2)
    for item in [
        "External user trial and feedback collection",
        "Frontend resource monitoring panel (CPU/memory/disk per agent)",
        "GitHub Actions CI with minimum 60% test coverage",
        "collector.py refactor (split 866-line file into focused modules)",
    ]:
        en_bullet(doc, item)

    en_hdr(doc, "8.3 Medium-Term", level=2)
    for item in [
        "PostgreSQL migration path for 50+ agent deployments",
        "Multi-language frontend (Chinese, English, Japanese)",
        "Plugin marketplace for third-party data source adapters",
    ]:
        en_bullet(doc, item)

    en_hdr(doc, "8.4 Explicitly Out of Scope", level=2)
    for item in [
        "Squad/team management (Multica covers this)",
        "Fully autonomous agent execution (safety risk — keep human-in-the-loop)",
        "Electron desktop app (web + CLI is sufficient)",
        "Multi-workspace management (achievable via multiple instances)",
    ]:
        en_bullet(doc, item)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 9: DEPLOYMENT
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "9. Deployment", level=1)

    en_hdr(doc, "9.1 Docker (Recommended)", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "# docker-compose.yml\n"
        "services:\n"
        "  myagentwatch:\n"
        "    build: .\n"
        "    ports:\n"
        '      - "10000:10000"\n'
        "    volumes:\n"
        "      - ~/.local/share/opencode:/data/opencode:ro\n"
        "      - ./config.yaml:/app/config.yaml\n"
        "      - ./data:/app/data\n"
        "    restart: unless-stopped"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    en_hdr(doc, "9.2 Local Development", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "cd myagentwatch\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate    # Windows\n"
        "source .venv/bin/activate   # macOS/Linux\n"
        "pip install -r requirements.txt\n"
        "python app.py\n"
        "# Open http://localhost:10000"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    en_hdr(doc, "9.3 CLI Installation", level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        "cd myagentwatch-cli\n"
        "pip install -e .\n"
        "myaw connect --server http://localhost:10000 --key myaw_xxx\n"
        "myaw daemon start"
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    page_break(doc)

    # ══════════════════════════════════════════════════════════
    # APPENDIX: QUICK REFERENCE
    # ══════════════════════════════════════════════════════════
    en_hdr(doc, "Appendix: Quick Reference", level=1)

    en_hdr(doc, "A.1 Tech Stack", level=2)
    en_table(doc,
        ["Layer", "Technology", "Rationale"],
        [
            ["Backend", "Python 3.12 + Flask 3.x", "Python is the AI ecosystem's lingua franca; Flask is lightweight and flexible"],
            ["Real-time Push", "Flask-SocketIO + SSE", "WebSocket for structured data; SSE for event streams; complementary channels"],
            ["Scheduler", "APScheduler", "Most mature in-process Python scheduler; no external Redis/Celery needed"],
            ["Config", "PyYAML + template engine", "Config-driven; industry templates (quant trading, web dev) for one-click setup"],
            ["System Monitor", "psutil + Windows fallback", "Cross-platform resource collection; PowerShell fallback on Windows"],
            ["Charts", "Chart.js (CDN)", "Zero build toolchain; CDN import and use immediately"],
            ["Topology", "dagre + d3.js (CDN)", "dagre layout + d3.js rendering for interactive DAG flow"],
            ["Storage", "SQLite3 × 2 + gzip archive", "Source DB read-only + aggregate DB read-write; single-file zero-ops"],
            ["Container", "Docker + docker-compose", "One-click deploy; volume mounts for persistence; read-only mounts protect source data"],
        ]
    )

    en_hdr(doc, "A.2 State Machine", level=2)
    en_table(doc,
        ["State", "Color", "Meaning"],
        [
            ["active", "Green", "Running normally with recent heartbeat or activity"],
            ["working", "Blue", "Explicitly executing a task"],
            ["idle", "Yellow", "Idle, awaiting first heartbeat, or process detected but no activity"],
            ["error", "Red", "Recent errors detected"],
            ["blocked", "Orange", "Persistent errors requiring human intervention"],
            ["offline", "Purple", "Disconnected or heartbeat lost"],
        ]
    )

    en_hdr(doc, "A.3 Alert Rules", level=2)
    en_table(doc,
        ["Rule", "Metric", "Threshold", "Level"],
        [
            ["agent_idle", "Time since last activity", "> 3600s", "warn"],
            ["high_cost", "Single session token cost", "> $5.00", "warn"],
            ["tool_failure_rate", "Failed / total tool calls", "> 20%", "critical"],
            ["cache_hit_rate_low", "cache_read / (read + write)", "< 30%", "info"],
        ]
    )

    en_hdr(doc, "A.4 Performance Optimizations", level=2)
    en_table(doc,
        ["Optimization", "Before", "After", "Impact"],
        [
            ["daily_stats aggregation", "Every 2s", "Once per day", "43,200× fewer writes"],
            ["Poll interval", "2s", "5s", "60% less read pressure"],
            ["Worker threads", "4", "8", "2× parallel throughput"],
            ["socketio.emit", "Synchronous (blocking HTTP)", "Background ThreadPoolExecutor", "HTTP latency from 200ms+ → < 20ms"],
            ["actual_last calculation", "heartbeat only", "max(heartbeat, activity_log, last_seen_time)", "Eliminated false offline for older agents"],
            ["Process detection", "None", "agent_processes → idle fallback", "No more false offline when process exists but heartbeat absent"],
        ]
    )

    # ── save ──────────────────────────────────────────────────
    out_path = Path.home() / "Desktop" / "MyAgentWatch-Project-Introduction-v6.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


# ═══════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    build_cn()
    build_en()
    print("All done.")
