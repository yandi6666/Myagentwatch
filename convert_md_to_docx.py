"""Convert MyAgentWatch markdown project introductions to .docx files."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

PROJECT_DIR = Path(__file__).resolve().parent

FILES = [
    ("MyAgentWatch-项目介绍.md", "MyAgentWatch-项目介绍.docx"),
    ("MyAgentWatch-Project-Introduction.md", "MyAgentWatch-Project-Introduction.docx"),
]


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            from lxml import etree
            element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        for attr, attr_val in val.items():
            element.set(qn(f'w:{attr}'), str(attr_val))


def add_styled_paragraph(doc, text, style=None, bold=False, font_size=None, color=None, alignment=None, space_after=None):
    """Add a paragraph with optional formatting."""
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if font_size:
            run.font.size = Pt(font_size)
        if color:
            run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def parse_inline(text, paragraph):
    """Parse inline markdown: bold, italic, code."""
    # Bold: **text**
    parts = []
    remaining = text
    while remaining:
        bold_match = re.match(r'(.*?)\*\*(.+?)\*\*', remaining)
        if bold_match:
            if bold_match.group(1):
                parts.append(('normal', bold_match.group(1)))
            parts.append(('bold', bold_match.group(2)))
            remaining = remaining[bold_match.end():]
        else:
            parts.append(('normal', remaining))
            break

    for ptype, content in parts:
        if ptype == 'bold':
            run = paragraph.add_run(content)
            run.bold = True
        else:
            # Handle inline code
            code_parts = re.split(r'`([^`]+)`', content)
            for i, cp in enumerate(code_parts):
                if i % 2 == 0:
                    paragraph.add_run(cp)
                else:
                    run = paragraph.add_run(cp)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.name = 'Calibri'
        if level == 1:
            h_font.size = Pt(22)
            h_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        elif level == 2:
            h_font.size = Pt(16)
            h_font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
        elif level == 3:
            h_font.size = Pt(13)
            h_font.color.rgb = RGBColor(0x16, 0x21, 0x3E)

    content = md_path.read_text(encoding='utf-8')

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                # Add shading
                from lxml import etree
                shading = etree.SubElement(p._p.get_or_add_pPr(), qn('w:shd'))
                shading.set(qn('w:fill'), 'F0F0F0')
                shading.set(qn('w:val'), 'clear')
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            from lxml import etree
            pPr = p._p.get_or_add_pPr()
            pBdr = etree.SubElement(pPr, qn('w:pBdr'))
            bottom = etree.SubElement(pBdr, qn('w:bottom'))
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            i += 1
            continue

        stripped = line.strip()

        # Table handling
        if stripped.startswith('|') and stripped.endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # Skip separator rows like |---|---|
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped[1:-1].split('|')]
            table_rows.append(cells)
            i += 1
            continue
        elif in_table:
            # End of table - render it
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for row_idx, row_data in enumerate(table_rows):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell = row.cells[col_idx]
                            cell.text = ''
                            p = cell.paragraphs[0]
                            parse_inline(cell_text, p)
                            # Header row styling
                            if row_idx == 0:
                                for run in p.runs:
                                    run.bold = True
                                from lxml import etree
                                shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
                                shading.set(qn('w:fill'), '0F3460')
                                shading.set(qn('w:val'), 'clear')
                                for run in p.runs:
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

                doc.add_paragraph()  # spacing after table
            table_rows = []
            in_table = False
            # Don't increment i - reprocess this line
            continue

        # Heading 1: ## with Chinese numbers or plain #
        h1_match = re.match(r'^#\s+(.+)', stripped)
        h2_match = re.match(r'^##\s+(.+)', stripped)
        h3_match = re.match(r'^###\s+(.+)', stripped)

        if h1_match:
            doc.add_heading(h1_match.group(1), level=1)
        elif h2_match:
            doc.add_heading(h2_match.group(1), level=2)
        elif h3_match:
            doc.add_heading(h3_match.group(1), level=3)

        # Blockquote
        elif stripped.startswith('> '):
            quote_text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(quote_text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Unordered list
        elif re.match(r'^[-*]\s+', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            parse_inline(text, p)

        # Ordered list
        elif re.match(r'^\d+\.\s+', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            p.clear()
            parse_inline(text, p)

        # Empty line
        elif stripped == '':
            # Skip consecutive empty lines
            pass

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            parse_inline(stripped, p)

        i += 1

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    doc.save(docx_path)
    print(f"  → {docx_path.name}")


def main():
    for md_name, docx_name in FILES:
        md_path = PROJECT_DIR / md_name
        docx_path = PROJECT_DIR / docx_name
        if not md_path.exists():
            print(f"SKIP: {md_name} not found")
            continue
        print(f"Converting: {md_name}")
        convert_md_to_docx(md_path, docx_path)
    print("Done.")


if __name__ == '__main__':
    main()
