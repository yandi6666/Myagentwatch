"""Generate MyAgentWatch presentation speech script."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def title(text, size=16, color=(0x0F, 0x34, 0x60)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p


def speak(text, bold=False):
    """Actual speech content."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    return p


def stage(text):
    """Stage direction / presentation tip."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    return p


def divider():
    p = doc.add_paragraph()
    run = p.add_run("─" * 50)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return p


# ═══════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("MyAgentWatch 项目答辩演讲稿")
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("面向：同学与老师  |  预计时长：25 分钟  |  2026 年 6 月")
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

stage("【使用说明】灰色斜体为演讲提示和舞台动作建议，正常字体为演讲稿正文。")
stage("【时间分配】开场 2min → 痛点 5min → 方案 8min → 架构 3min → 竞品 2min → 现状与展望 3min → 致谢 2min")
stage("PPT 页面切换标记为 ▶ 符号，建议每页停留 1-3 分钟。")

doc.add_page_break()

# ═══════════════════════════════════════════
# PART 1: OPENING
# ═══════════════════════════════════════════
title("第一部分：开场（2 分钟）")
divider()

stage("【PPT 第 1 页：封面】站定，微笑，眼神扫过全场。")

speak("各位老师、各位同学，大家下午好。")
speak("")
speak("我今天答辩的项目叫做 MyAgentWatch。")
speak("")
speak("在开始之前，我想先问大家一个问题——在座有多少同学平时在用 AI 编程助手？Claude Code、GitHub Copilot、或者 Cursor？")
speak("")
stage("【等待举手，互动】")
speak("好，不少同学都在用。那我再问一个问题——你们知道 AI 帮你写代码的时候，它到底在想什么吗？它调用了哪些工具？读了你哪些文件？它每次帮你写代码花了多少钱？")
speak("")
speak("如果你不知道——没关系，因为没有人知道。这就是 MyAgentWatch 要解决的问题。")
speak("")
speak("一句话介绍我的项目：让你像管理微服务一样管理你的 AI Agent 团队——实时看清每个 Agent 在想什么、做什么、花了多少钱。")

stage("【PPT 翻页】")

# ═══════════════════════════════════════════
# PART 2: PAIN POINTS
# ═══════════════════════════════════════════
title("第二部分：为什么需要 MyAgentWatch？（5 分钟）")
divider()

stage("【PPT 第 2 页：痛点总览 — 三大黑盒】")

speak("2025 年以来，AI Agent 从实验品变成了生产工具。越来越多的团队把 Agent 当成虚拟团队成员——让它做需求分析、架构设计、代码生成、Bug 修复。")
speak("")
speak("但是，Agent 有一个致命的问题：它是黑盒。具体来说，三个黑盒——")

speak("第一个黑盒：通信黑盒。你看不到 Agent 之间在说什么。", bold=True)
speak("")
speak("我给大家讲一个真实的例子——也是我自己每天都在经历的。")

stage("【PPT 第 3 页：Before — OpenClaw 调用 OpenCode 的流程】指着流程图讲解。")

speak("OpenClaw 是一个通用 AI Agent 编排框架，你可以把它理解成 AI 团队里的「项目经理」。OpenCode 是一个专注于终端编程的 AI Agent，相当于「程序员」。")
speak("")
speak("典型的工作流是这样的——用户对 OpenClaw 说：「帮我写一个 Flask REST API」。OpenClaw 内部通过一个叫 ACP 的底层协议，静默启动了一个 OpenCode 子进程，把编码任务委托给它。OpenCode 开始读文件、写代码、跑测试——整个过程在子进程中静默执行。最后 OpenCode 把结果返回给 OpenClaw，OpenClaw 用自己的话转述给用户。")
speak("")
speak("问题在哪？用户从头到尾只看到 OpenClaw 说「处理好了」或「遇到了一些问题」。中间的委托指令是什么？OpenCode 有没有误解？报了什么错？OpenClaw 有没有篡改输出？全部不可见。")
speak("")
speak("这就像两个同事关上门在小房间讨论，你只听到最终结论。中间吵了什么、改了什么、为什么改，你全都不知道。")

stage("【停顿 2 秒，加重语气】")

speak("更糟的是——如果 OpenCode 陷入死循环呢？")

stage("【PPT 第 4 页：死循环代价账单】")

speak("假设 OpenCode 写代码时遇到了一个它不理解的 import 错误。它第一轮写入有 bug 的代码，测试失败了。第二轮它说「我理解了，是 import 路径问题」，修改后还是失败。第三轮到第八轮，它反复读同一个文件，反复猜测原因，反复失败。每次循环消耗八千到一万五千 input tokens、两千到五千 output tokens。第九轮，它的思考开始退化——上下文窗口被无效信息填满了——它说「我建议回滚所有改动重来」，然后开始新一轮循环。")
speak("")
speak("三十分钟后用户打开 OpenClaw，看到一句轻描淡写的「任务处理遇到一些困难」。但实际账单是——三十八万 input tokens、九万五 output tokens，按 Claude Sonnet 定价约 1.71 美元。源文件被无效修改了十二次。用户完全不知道这三十分钟里发生了什么。")
speak("")
speak("如果这发生在凌晨三点的自动化任务中——第二天团队发现任务没完成、Token 被刷掉了十五美元以上、代码仓库多了一堆需要回滚的提交。")
speak("")
speak("这不是假设。Claude Opus 4 单次请求可以消耗三万两千 output tokens，约 0.48 美元。不加监控，一个循环十轮的 Agent 足以在一小时内烧掉五到十五美元。五个 Agent 并行的量化交易团队，一天的损失可能高达两百美元。")

stage("【PPT 第 5 页：After — MyAgentWatch 如何解决】")

speak("而有了 MyAgentWatch 之后呢？OpenClaw 在群聊里 @OpenCode，发送委托消息——这条消息以紫色 Handoff 卡片实时显示，清晰标注「谁委托谁、委托什么、什么状态」。OpenCode 在群聊里回复自己的执行进展——蓝色思考气泡、橙色工具调用框、绿色输出结果。SSE 事件流实时推送每一步操作：思考了 2.3 秒、读取了 app.py、写入了 routes 文件、运行了 pytest、全部通过。")
speak("")
speak("全程可追溯——谁委托了谁、委托了什么、怎么做的、每步花了多少 Token 和多少时间——全部有记录。当 OpenCode 第三次读取同一个文件时，告警引擎触发 blocked 状态，用户实时收到通知，在群聊中直接打断循环。Token 仪表盘实时显示消耗曲线陡然上升——一眼就能发现异常。")

speak("第二个黑盒：成本黑盒。团队不知道钱花在哪。", bold=True)
speak("")
speak("一个中型团队每天消耗数百万 Token。Claude Opus 和 DeepSeek V4 的价差可以达到五十倍以上。不同的 Agent 类型消耗差异巨大——plan 层通常远高于 build 层。不同时段消耗也不均衡。但没有任何工具能告诉团队「今天花了多少钱、哪个 Agent 最费钱、缓存有没有生效、哪些模型还没定价」。")

speak("第三个黑盒：运维黑盒。Agent 停了都没人知道。", bold=True)
speak("")
speak("Agent 可能因为 API 限流、Token 耗尽、内存溢出、数据源断开这些原因挂掉。传统 APM 工具看不懂 AI Agent 的专属指标——心跳、模型状态、工具调用成功率、缓存效率。Agent 凌晨三点宕掉，团队第二天才发现，浪费大量计算资源和时间窗口。")

stage("【PPT 翻页，准备进入方案部分】")

# ═══════════════════════════════════════════
# PART 3: WHAT IS MYAGENTWATCH
# ═══════════════════════════════════════════
title("第三部分：MyAgentWatch 是什么？（8 分钟）")
divider()

stage("【PPT 第 6 页：项目一句话 + 核心数据】")

speak("MyAgentWatch 是一个开源的、自托管的 AI Agent 实时监控与协作平台。核心定位是 Agent Observability——Agent 可观测性平台。")
speak("")
speak("几个关键数字：支持八家 AI 厂商三十七个模型的定价；四十多个 REST API 端点加七个 WebSocket 事件加 SSE 实时流；十一个 CLI 命令；采集周期五秒；SQLite 单文件数据库，零运维。")
speak("")

stage("【PPT 第 7 页：核心功能总览】")

speak("项目有六大核心功能模块——")
speak("")
speak("第一个，Token 用量仪表盘——省钱的第一道防线。", bold=True)
speak("五级 Token 拆解——input、output、reasoning、cache_read、cache_write，比业界常见的三级拆解更精细。五种查询维度——总览、按 Agent、按模型、按小时趋势、未映射模型诊断。八家厂商三十七个模型的定价覆盖国内外主流厂商。前端的柱状图、双数据表、趋势折线图、未映射诊断面板，让你一目了然。")
speak("")

stage("【PPT 第 8 页：群聊系统】")

speak("第二个，Agent 群聊系统——核心差异化武器。", bold=True)
speak("这是六月一日全新重做的微信风格三栏布局。左栏是会话列表——圆形头像、在线绿点、最后消息预览、红点未读计数。中栏是消息流——人类蓝色气泡靠右、Agent 灰色气泡靠左，支持四种富媒体消息：tool_call 橙色边框卡片、handoff 紫色边框卡片、share 绿色边框卡片、以及普通灰色气泡。右栏是 Agent 详情卡片和通讯录——按分组显示、在线绿点标记。")
speak("")
speak("我特别想强调一点——群聊系统的设计目标不是让 Agent「社交」，而是让 Agent 之间的通信可视化。你可以指定 Agent 之间必须用中文还是英文对话。Agent 之间的每一次委托、每一次交接，都变成了一条可追溯的群聊消息。就像看 Slack 频道里的讨论记录一样自然。")

stage("【PPT 第 9 页：SSE 事件流 + 告警 + CLI】")

speak("第三个，SSE 实时事件流——Agent 的「直播画面」。", bold=True)
speak("四种事件类型用四种颜色区分：thinking 蓝色、tool_call 橙色、response 绿色、handoff 紫色。页面加载即启动，不等待用户切换 tab——这是根据用户反馈修复的关键 UX 问题。支持按 Agent 筛选、按类别筛选、星标收藏、仅错误模式、暂停恢复和一键清空。点击任何事件展开完整文本内容。")
speak("")
speak("第四个，告警规则引擎——主动发现问题。四条默认规则覆盖 Agent 空闲、高成本、工具失败率、缓存命中率。全部通过 config.yaml 配置，支持热修改和自定义扩展。告警触发后写入数据库、WebSocket 推送、前端 Toast 通知——三层通知链。")
speak("")
speak("第五个，CLI 命令行客户端。十一个命令，纯 Python 标准库实现，零外部依赖。connect 连接、status 仪表盘、tokens 用量、chat 群聊、heartbeat 心跳守护模式——让 Agent 能自己管理自己的状态。")
speak("")
speak("第六个，六状态 Agent 状态机——active、working、idle、error、blocked、offline。综合 heartbeat、activity_log、last_seen_time 三大信号源判定。比 Multica 多 working 状态，比传统工具精细得多。")

stage("【PPT 翻页，过渡到架构部分】")

# ═══════════════════════════════════════════
# PART 4: ARCHITECTURE
# ═══════════════════════════════════════════
title("第四部分：技术架构（3 分钟）")
divider()

stage("【PPT 第 10 页：四层架构图】")

speak("技术架构上，MyAgentWatch 采用经典的四层模型——展示层、API 层、业务层、数据层。")
speak("")
speak("最核心的设计模式是适配器模式。所有数据源实现统一的 SourceInterface 抽象基类，通过装饰器注册。新增数据源——三步。第一，新建适配器文件，实现三个方法；第二，用 @register_source 注册；第三，config.yaml 加一行 type 声明。核心代码一行不改。这就是「对扩展开放，对修改封闭」。")
speak("")
speak("技术选型上我说三个关键决策——")
speak("")
speak("第一，为什么选 SQLite 而不是 PostgreSQL？因为零部署——一个文件即完整数据库，备份就是 cp。三到十个 Agent 的场景下 SQLite WAL 模式完全够用。五十个 Agent 以上平滑迁移 PostgreSQL——所有 SQL 都是标准 SQL，无 ORM 绑定。")
speak("")
speak("第二，为什么同时用 WebSocket 和 SSE？因为不同数据需要不同的推送方式。WebSocket 推结构化小消息——状态变更、统计数据，适合双向通信和房间机制。SSE 推非结构化大文本——thinking 可能数千字，适合单向流式推送。两个通道各司其职，不重叠。")
speak("")
speak("第三，为什么前端零构建工具链？当前阶段快速迭代的务实选择——修改 JS 立即刷新。当代码量增长后会引入 Vite 加 ES modules，但坚持 vanilla JS，不引入 React 或 Vue 框架。保持轻量是原则。")

stage("【PPT 翻页，过渡到竞品对比】")

# ═══════════════════════════════════════════
# PART 5: COMPETITIVE ANALYSIS
# ═══════════════════════════════════════════
title("第五部分：竞品对比与差异化（2 分钟）")
divider()

stage("【PPT 第 11 页：竞品对比矩阵】")

speak("和 Multica 的对比——这是目前市场上功能最完整的 Agent 管理平台之一。Go 加 PostgreSQL 加 Redis，技术栈很重。重点看差异——")
speak("")
speak("MyAgentWatch 有但 Multica 没有的：系统资源监控——CPU、内存、磁盘。Agent 群聊社交——Multica 完全没有这个方向。告警规则引擎。SSE 实时事件流。")
speak("")
speak("Multica 有但 MyAgentWatch 还缺的：完整的任务生命周期管理。这个在路线图里——六月十日前完成。")
speak("")
speak("一句话总结：Multica 管任务，MyAgentWatch 管 Agent。两者不是替代关系，是互补关系。")
speak("")
speak("和通用工具的对比——为什么不直接用 Grafana 加 Langfuse 加 Datadog 的组合？两个致命缺陷：数据孤岛——Grafana 看到 CPU 正常、Langfuse 看到 API 调用成功，但你连不成一条线说「OpenClaw 在第五轮把任务 handoff 给了 OpenCode」。部署成本——Langfuse 需要 PostgreSQL 加 ClickHouse 加 Redis。和一个核心缺失——这些通用工具没有 Agent 状态机、没有 handoff 追踪、没有 Token 按 Agent 拆解。")

stage("【PPT 翻页，过渡到现状与展望】")

# ═══════════════════════════════════════════
# PART 6: CURRENT STATUS & ROADMAP
# ═══════════════════════════════════════════
title("第六部分：现状与展望（3 分钟）")
divider()

stage("【PPT 第 12 页：项目现状】")

speak("我想诚实地向大家汇报项目当前的阶段——不夸大、不回避。")
speak("")
speak("MyAgentWatch 目前是 v2.1 版本，一个功能完整的 Alpha 产品。在我的本地环境持续运行了三周，监控五个 Agent——Claude Code 的 plan、build、explore，加上 OpenCode 实例。每天产生两百到五百条 Token 记录、五十到一百条对话 Turn。")
speak("")
speak("已经验证的功能：Token 成本监控准确追踪了五到六月的 Token 消费；SSE 事件流能实时展示 Agent 的 thinking 和 tool_call；群聊能展示紫色的 handoff 卡片。")
speak("")
speak("尚待验证的：告警引擎的误报漏报率，需要更多异常场景的测试。大规模 Agent 的并发性能，目前只有五个 Agent 的体验数据。外部用户的反馈——目前只有我自己在用。")

stage("【PPT 第 13 页：路线图】")

speak("六月的三个优先目标——")
speak("")
speak("第一优先级，Agent Backend 抽象层，两到三天。统一接口接入多种 Agent CLI——Claude Code、OpenCode、以及未来的更多工具。这是底座。")
speak("")
speak("第二优先级，基础任务系统，五到七天。补上 Multica 有但我们缺的能力——从 queued 到 assigned 到 running 到 completed 或 failed 的完整任务生命周期。")
speak("")
speak("第三优先级，Agent 群聊增强，三到四天。增加 @提及、消息引用回复、表情回应、群公告——让群聊从「能看」变成「能用」。")
speak("")
speak("明确不做的事情：Autopilot 自动执行——让监控系统自动执行 Agent 任务是危险的，保持「监控不干预」的定位。Electron 桌面客户端——Web 加 CLI 已经够了。Squad 小队管理——Multica 已经做好的不重复造轮子。")

stage("【PPT 第 14 页：北极星指标】")

speak("最后，我想分享项目的北极星指标——判断它是否成功的三个标准：")
speak("")
speak("第一，Token 成本可追溯率——多少百分比的 Token 消耗能精确归属到具体的 Agent、Session、任务。目标大于百分之九十五。")
speak("")
speak("第二，故障发现时间——从 Agent 出现异常到用户收到通知的延迟。目标小于三十秒。")
speak("")
speak("第三，Agent 间通信的透明化覆盖率——多少次 Agent handoff 被记录并可视化。目标百分之百。")
speak("")
speak("这三个指标直接对应三大痛点：省钱、稳定性、透明化。")

# ═══════════════════════════════════════════
# PART 7: CLOSING
# ═══════════════════════════════════════════
title("第七部分：总结与致谢（2 分钟）")
divider()

stage("【PPT 第 15 页：总结金句】站定，面向观众。")

speak("让我用三句话总结这个项目——")
speak("")
speak("第一句，问题——AI Agent 正在成为开发团队的标配，但我们看不到它们在做什么、花了多少钱、什么时候出错。这是一个真实的、正在加速的痛点。")
speak("")
speak("第二句，方案——MyAgentWatch 把 Agent 之间的通信从黑盒变成白盒，让每一次委托、每一次工具调用、每一分钱都有迹可循。")
speak("")
speak("第三句，愿景——如果我们能让 Agent 像微服务一样可观测、可管理、可审计，那整个 AI Agent 生态的可靠性和效率都会上一个台阶。")
speak("")
speak("我最想让大家记住的一句话是——")
speak("")
speak("MyAgentWatch 让 AI Agent 不再是一个黑盒。就像 Kubernetes 让微服务可管理一样，MyAgentWatch 让 AI Agent 可观测。", bold=True)

stage("【停顿 3 秒】")

speak("谢谢各位老师和同学。以上就是我的答辩。下面欢迎大家提问。")

stage("【鞠躬，等待 Q&A】")

doc.add_page_break()

# ═══════════════════════════════════════════
# APPENDIX: Q&A PREPARATION
# ═══════════════════════════════════════════
title("附录：答辩 Q&A 准备（常见刁钻问题速答）", size=14, color=(0x99, 0x99, 0x99))
divider()

qas = [
    ("Q1：这不是只是把日志换了个好看的壳吗？",
     "如果只是换壳，能做到跨 Agent 的通信可视化和 Token 到 Agent 的成本归因吗？日志是碎片化的——不同 CLI 格式不同、不同 Agent 的日志互不关联、没有实时性。MyAgentWatch 做的是把这些碎片拼成完整的故事线。"),
    ("Q2：Agent 社交是刚需还是伪需求？",
     "对个人开发者不是刚需，对多 Agent 团队是刚需——但刚需的不是「聊天」，而是「Agent 间通信的可视化和审计」。群聊只是这个能力的 UI 表达。不喜欢 UI 可以用 CLI 或 API。"),
    ("Q3：如果只留一个功能，留什么？",
     "Token 用量仪表盘。因为它是唯一直接关系到「钱」的功能——所有用户都关心花了多少钱。"),
    ("Q4：状态机误判率多高？有没有真实验证？",
     "量化数据尚未收集——需要真实生产环境的标注数据。这是必须诚实的回答。但防震荡机制已设计好：多信号源交叉验证、2 倍超时 offline 判定、status_since 追溯。"),
    ("Q5：测试覆盖率多少？有 CI/CD 吗？",
     "当前不足 30%，无 CI/CD。这是最大的技术债。已在 6 月路线图列为 P0。核心模块（状态机、定价计算）有测试，数据源适配器和前端缺乏系统测试。"),
    ("Q6：如果大厂下场做 Agent observability，怎么办？",
     "三层护城河：厂商中立的 Agent 通信总线、跨 8 厂 37 模型的统一定价模型、极简部署。大厂工具只覆盖自家模型，MyAgentWatch 是多厂商统一面。速度是关键。"),
    ("Q7：删掉所有 UI 之后，还剩什么？",
     "REST API 40+ 端点 + SSE 事件流 + CLI 11 命令 + collector 引擎 + 状态机 + 告警引擎。完整后端可驱动任何前端。架构是前后端分离的——前端可替换。"),
    ("Q8：为什么不现在就商业化？",
     "v2.1 是 Alpha 阶段，需要在 3-5 个真实团队中验证产品价值。先验证再商业化。远期路径：开源核心免费 + 企业版付费 + SaaS 托管。至少 2027 年以后。"),
]

for q, a in qas:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    p = doc.add_paragraph()
    run = p.add_run(a)
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.add_paragraph()

# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "MyAgentWatch-答辩演讲稿.docx")
doc.save(output_path)
print(f"演讲稿已生成: {output_path}")
print("正文约 4500 字，预计演讲时长 20-25 分钟 + Q&A")
