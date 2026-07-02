"""Extract v7 docx structure for reference."""
from docx import Document

doc = Document(r"C:\Users\天宇\Desktop\MyAgentWatch-项目介绍-v7.docx")

with open("_v7_extract.txt", "w", encoding="utf-8") as f:
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "None"
        text = para.text
        if text.strip():
            f.write(f"[{style}] {text}\n\n")

    f.write("\n=== TABLES ===\n")
    for ti, table in enumerate(doc.tables):
        f.write(f"\nTable {ti}: {len(table.rows)} rows x {len(table.columns)} cols\n")
        for ri, row in enumerate(table.rows):
            cells = [cell.text[:100] for cell in row.cells]
            f.write(f"  Row {ri}: {cells}\n")

print("Done. See _v7_extract.txt")
