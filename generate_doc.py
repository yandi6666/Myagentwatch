"""Generate MyAgentWatch project introduction Word document for PPT conversion."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from faq_section import add_faq_chapter

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return h


def para(text, bold=False, size=None, color=None, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment is not None:
        p.alignment = alignment
    return p


def bullet(text, indent=0):
    """Add a bullet point with optional indent level."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1 + indent)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row_obj in t.rows:
                row_obj.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def callout(text):
    """Highlight box for key numbers or quotes."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x34, 0x60)
    return p


# ═══════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
para("MyAgentWatch", bold=True, size=40, color=(0x0F, 0x34, 0x60), alignment=WD_ALIGN_PARAGRAPH.CENTER)
para("多 Agent 实时监控仪表盘", bold=True, size=20, color=(0x55, 0x55, 0x55), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
para("打破 AI Agent 黑盒 —— 实时追踪全链路行为与成本", size=14, color=(0x77, 0x77, 0x77), alignment=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
para("版本 2.1    |    2026 年 6 月", size=11, color=(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)
para("Python 3.12  ·  Flask  ·  SocketIO  ·  SQLite  ·  Chart.js", size=10, color=(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)
para("面向：量化交易 AI 团队  |  通用 AI 开发团队", size=10, color=(0x99, 0x99, 0x99), alignment=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 一、项目背景与痛点
# ═══════════════════════════════════════════════════════════════
heading("一、项目背景：为什么需要 MyAgentWatch？", level=1)

heading("1.1 行业趋势：AI Agent 正在成为开发团队的标准配置", level=2)
para(
    "2025 年以来，以 Claude Code、OpenCode、Codex、Cursor、Gemini CLI 为代表的 AI 编程 Agent 迅速普及。"
    "越来越多的团队将 Agent 深度集成到日常开发流程中——代码生成、需求分析、架构设计、测试编写、Bug 修复。"
    "在量化交易等高频决策场景，AI Agent 还承担策略研究、因子挖掘、回测分析等关键任务。"
    "Agent 已经从「辅助工具」演变为「虚拟团队成员」，其运行状态、决策质量和资源消耗直接影响项目进度和成本。"
)

heading("1.2 三大核心痛点", level=2)

para("痛点一：Agent 间通信是黑盒——你看不到 Agent 之间说了什么", bold=True)

para(
    "以一个真实的多 Agent 协作场景为例：OpenClaw 是一款通用 AI Agent 编排框架（类似 AI 团队的「项目经理」），"
    "OpenCode 是一款专注于终端编程的 AI Agent（类似 AI 团队的「程序员」）。典型工作流是：用户给 OpenClaw 下达任务 "
    "→ OpenClaw 分析后将编码部分委托给 OpenCode → OpenCode 完成编码把结果返回给 OpenClaw → OpenClaw 汇总呈现给用户。"
)

para("▎ Before（使用 MyAgentWatch 之前）—— 用户视角：", bold=True)
bullet("● 用户在 OpenClaw 的聊天界面中说：「帮我写一个 Flask REST API」")
bullet("● OpenClaw 内部通过 ACP（Agent Communication Protocol，JSON-RPC 2.0 over NDJSON 子进程协议）启动了一个 OpenCode 子进程")
bullet("● OpenClaw 向 OpenCode 发送了编码指令，但这是底层 IPC 通信——用户完全看不到指令内容")
bullet("● OpenCode 开始读文件、写代码、跑测试，整个过程在终端子进程中静默执行")
bullet("● OpenCode 完成后返回结果给 OpenClaw，OpenClaw 用自己的语言复述结果给用户")
bullet("● 用户的困惑：OpenClaw 到底让 OpenCode 做了什么？指令是否准确？OpenCode 有没有误解？OpenCode 报了什么错？OpenClaw 有没有篡改/简化 OpenCode 的输出？全部不可见")
bullet("● 更糟的是：如果 OpenCode 陷入死循环（反复读同一个文件）、调用了错误的工具、或者返回了错误的结果，用户在 OpenClaw 界面完全感知不到——只能看到 OpenClaw 说「处理好了」或「遇到了一些问题」")

para("\n▎ OpenCode 陷入死循环的真实代价：", bold=True)
para(
    "假设 OpenCode 在编写代码时遇到了一个它不理解的 import 错误。在没有监控的情况下，一个典型的死循环场景是这样的："
)
bullet("● 第 1 轮：OpenCode 写入了一段有 bug 的代码 → 运行测试 → 失败 → 「让我看看错误信息」→ 读取错误日志")
bullet("● 第 2 轮：OpenCode 「我理解了，是 import 路径问题」→ 修改代码 → 运行测试 → 又失败 → 「不对，让我重新看看」→ 再次读取同一个文件")
bullet("● 第 3-8 轮：OpenCode 反复读同一个文件 → 反复猜测原因 → 反复修改 → 反复失败。每次循环消耗 8,000-15,000 input tokens + 2,000-5,000 output tokens")
bullet("● 第 9 轮：OpenCode 的思考开始退化（上下文窗口被无效信息填满）→ 「我建议回滚所有改动重来」→ 又开始新一轮循环")
bullet("● 最终结果：30 分钟后用户打开 OpenClaw，看到一条轻描淡写的「任务处理遇到一些困难，正在重试」。但实际上：\n"
       "    · 已消耗 380,000 input tokens（成本约 $1.14，按 Claude Sonnet 定价）\n"
       "    · 已消耗 95,000 output tokens（成本约 $0.57）\n"
       "    · 总共浪费 $1.71，耗时 30 分钟\n"
       "    · 源文件被反复修改了 12 次，git diff 显示大量无效改动\n"
       "    · 用户完全不知道这 30 分钟里发生了什么")
bullet("● 如果这个场景发生在凌晨 3 点的自动化任务中，团队第二天早上才会发现：任务没完成、Token 被刷掉了 $15+、代码仓库多了一堆需要回滚的提交")

para("\n这不是假设。在实际使用中，Claude Opus 4 单次请求可消耗高达 32,000 output tokens（约 $0.48），如果不加监控，一个循环 10 轮的 Agent 足以在一小时内烧掉 $5-15。对于同时运行 5 个 Agent 的量化交易团队，一天的死循环损失可能高达 $200。")

para("\n▎ 而有了 MyAgentWatch：", bold=True)
bullet("● 当 OpenCode 第 3 次读取同一个文件时，告警规则 tool_failure_rate 或 agent_idle（反复执行无进展）触发橙色 blocked 状态")
bullet("● 用户实时收到通知：「OpenCode 在过去 5 分钟内读取了同一个文件 3 次，疑似死循环」")
bullet("● 用户在群聊中直接 @OpenCode 发送：「停止当前任务，换一种思路」——人工介入，打破循环")
bullet("● 或者用户查看 SSE 事件流，发现 OpenCode 的 thinking 内容越来越长、越来越混乱（上下文污染），直接通过 CLI 发送 heartbeat --status blocked 强制挂起 Agent")
bullet("● Token 仪表盘实时显示 OpenCode 的消耗曲线陡然上升，用户一眼就能发现异常")

para("\n这就是透明化的力量：不是阻止 Agent 犯错（Agent 必然会犯错），而是在犯错的第一时间让你知道，并给你干预的手段。")

para("\n这就像两个同事关上门在小房间里讨论，你只听到了最终「我们决定这样」的结论——中间的争论、纠错、方案变更你全都不知道。对于量化交易等关键场景，这种黑盒是不可接受的：你不知道一个交易策略是经过充分验证的还是草率决定的。")

para("\n▎ After（使用 MyAgentWatch 之后）—— 用户视角：", bold=True)
bullet("● 用户同样对 OpenClaw 说：「帮我写一个 Flask REST API」")
bullet("● OpenClaw 通过 MyAgentWatch 的群聊系统，在群聊中 @OpenCode 并发送委托消息：「请帮我创建一个 Flask REST API，包含 GET /users 和 POST /users 两个端点，使用 SQLite 存储，代码风格遵循 PEP 8。」")
bullet("● 这条委托消息以紫色 Handoff 卡片的形式实时显示在群聊消息流中，清晰标注「OpenClaw → OpenCode / subagent_type: coding / 状态: running」")
bullet("● OpenCode 收到任务后，在群聊中回复自己的执行进展——thinking 消息（蓝色气泡·思考过程，如「先检查项目结构，确定在哪里添加路由文件」）、tool_call 消息（橙色框·工具调用，如「Bash: ls app/routes/」→ exit_code: 0）、response 消息（绿色气泡·代码输出结果）")
bullet("● 同时，SSE 事件流实时推送 OpenCode 的每一步操作：思考了 2.3 秒 → 调用了 read 工具读取 app.py → 调用了 write 工具写入 routes/users.py → 调用了 bash 工具运行 pytest → 测试全部通过 → 返回结果给 OpenClaw")
bullet("● OpenClaw 收到结果后，在群聊中用自己的话总结回复用户，同时把 OpenCode 的原始输出以绿色 Share 卡片形式分享到群聊")
bullet("● 整个过程完整记录：谁委托了谁、委托了什么、子 Agent 怎么做的、每一步的 Token 消耗和耗时、最终结果是什么——全部可追溯、可审计")

para("\n这就是 MyAgentWatch 群聊系统作为「Agent 间通信总线」的核心价值：让 Agent 之间的每一次对话、每一次委托、每一次交接都透明可见。团队可以在事后回放任意一次多 Agent 协作的全过程，就像查看 Slack 频道里的讨论记录一样自然。")

heading("1.3 MyAgentWatch 的答案：Agent 通信基础设施", level=2)

para(
    "MyAgentWatch 不仅是一个监控仪表盘，更是一套 Agent 之间的通信基础设施。它的群聊系统实际上充当了 Agent 间通信的统一消息总线：\n\n"
    "沟通语言可指定 —— Agent 之间的所有通信必须通过群聊消息，用户可以要求 Agent 使用中文或英文进行对话。"
    "比如在量化交易场景，用户要求「所有 Agent 间通信必须使用中文」，则 OpenClaw 给 OpenCode 的委托指令必须是中文，"
    "OpenCode 的回复也必须是中文。这让非技术团队成员也能看懂 Agent 之间在讨论什么。\n\n"
    "消息类型区分 —— 普通消息（灰色气泡）、工具调用（橙色边框）、Agent 交接（紫色边框）、成果分享（绿色边框），"
    "不同颜色和边框样式让用户一眼就能分辨消息性质，快速定位关键信息。\n\n"
    "完全可追溯 —— 所有 Agent 间通信持久化存储在 conversation_turns + turn_content + agent_handoffs 三张表中，"
    "天然键 natural_key 去重保证数据完整。用户可以按 trace_id 回放任意一次多 Agent 协作的完整链路。\n\n"
    "Token 成本归属 —— 每条委托消息关联具体的 turn_id，Token 消耗精确归属到发起任务的 Agent。"
    "你可以明确知道「OpenClaw 委托 OpenCode 执行的任务，消耗了 45,000 input tokens + 12,000 output tokens，成本 $0.38」。"
)

heading("1.4 三大核心痛点（续）", level=2)

para("痛点二：Token 成本失控——团队不知道钱花在哪", bold=True)
para(
    "一个中型开发团队每天可能消耗数百万 Token。不同模型（Claude Opus vs DeepSeek V4 vs GLM-5）价差可达 50 倍以上。"
    "不同 Agent 类型消耗差异巨大（plan 层通常消耗远高于 build 层）。不同时段（业务高峰期 vs 深夜）消耗不均衡。"
    "没有任何工具能告诉团队：「今天花了多少钱？哪个 Agent 最费钱？缓存命中率是多少？有没有模型没用定价？」"
)

para("痛点三：缺乏运维保障——Agent 停了都没人知道", bold=True)
para(
    "Agent 进程可能因为各种原因挂掉：API 限流、Token 耗尽、内存溢出、数据源连接断开。"
    "传统 APM 工具（如 Grafana、Datadog）无法感知 AI Agent 的专属指标：心跳、模型状态、工具调用成功率、缓存效率。"
    "当 Agent 在凌晨 3 点自动化任务中宕掉，团队要等到第二天才发现，浪费大量计算资源和时间窗口。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 二、项目概述
# ═══════════════════════════════════════════════════════════════
heading("二、项目概述", level=1)

heading("2.1 MyAgentWatch 是什么？", level=2)
para(
    "MyAgentWatch 是一款开源的、自托管的 AI Agent 实时监控与协作平台。它以适配器模式自动接入各类 AI Agent CLI "
    "（Claude Code、OpenCode 及兼容工具），通过采集 SQLite 数据库和日志文件，提供从 Token 消耗、工具调用、"
    "会话链路到系统资源的全维度可视化仪表盘。同时内置 Agent 群聊社交系统、SSE 实时事件流、告警规则引擎和 CLI 命令行客户端，"
    "让团队能够透明化地管理 AI 工作负载。"
)

heading("2.2 核心价值主张", level=2)

callout("「让你像管理微服务一样管理你的 AI Agent 团队」")

bullet("● 黑盒透明化 —— 自动发现并实时展示每个 Agent 的思考过程、工具调用和输出结果")
bullet("● 全维度成本管控 —— 8 家厂商 37 个模型定价，按 Agent / 模型 / 小时 / 会话四级拆解")
bullet("● 实时协作 —— Agent 之间自由聊天、交接任务、分享成果，微信风格群聊界面")
bullet("● 生产级运维 —— CPU / 内存 / 磁盘监控 + 5 状态机 + 告警规则引擎 + 日志归档")
bullet("● 零配置启动 —— 无 config.yaml 时自动扫描默认路径，自动发现所有 Agent")
bullet("● 极低部署成本 —— SQLite 单文件数据库，Docker 一键部署，不依赖外部服务")

heading("2.3 关键数据", level=2)

callout("支持的 AI 厂商：8 家（OpenAI / Anthropic / Google / DeepSeek / 阿里 / 月之暗面 / 智谱 / 字节）")
callout("定价模型覆盖：37 个模型")
callout("API 端点：40+ 个 REST 端点 + 7 个 WebSocket 事件 + SSE 实时流")
callout("CLI 命令：11 个")
callout("采集周期：5 秒（可配置）")
callout("数据库：SQLite 单文件（零运维）")

heading("2.4 适用场景", level=2)
bullet("● 量化交易团队 —— 监控多个 AI Agent 并行执行策略研究和因子挖掘任务")
bullet("● 软件开发团队 —— 追踪 Claude Code / OpenCode Agent 的编码活动和 Token 成本")
bullet("● AI 初创公司 —— 向客户展示 AI Agent 运行透明度和成本控制能力")
bullet("● 个人开发者 —— 在本地监控自己的 AI 编程助手，了解使用习惯和花费")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 三、技术栈与选型理由
# ═══════════════════════════════════════════════════════════════
heading("三、技术栈与选型理由", level=1)

heading("3.1 整体技术架构", level=2)

table(
    ["层次", "技术选型", "为什么选它"],
    [
        ["后端框架", "Python 3.12 + Flask 3.x",
         "Python 是 AI 生态的 lingua franca；Flask 轻量灵活，适合中小规模项目快速迭代；3.12 支持最新语言特性"],
        ["实时推送", "Flask-SocketIO + SSE",
         "WebSocket 承载结构化数据推送（增量 2s + 全量 10s）；SSE 承载事件流（单向大文本）；双通道各司其职"],
        ["定时调度", "APScheduler",
         "Python 生态最成熟的进程内调度器，支持 cron 和 interval 两种模式，无需外部 Redis/Celery 依赖"],
        ["配置管理", "PyYAML + 模板引擎",
         "config.yaml 配置驱动，支持行业模板（量化交易 / Web 开发）一键切换告警规则和 Agent 分组"],
        ["系统监控", "psutil",
         "跨平台系统资源采集（Windows/Linux/macOS），CPU/内存/磁盘/网络全部覆盖，开销极小"],
        ["前端图表", "Chart.js (CDN)",
         "零构建工具链，CDN 引入即用；柱状图 + 折线图满足所有可视化需求"],
        ["拓扑可视化", "dagre + d3.js (CDN)",
         "dagre 布局算法 + d3.js 渲染引擎，生成 DAG 交互式 Agent 流程图"],
        ["数据存储", "SQLite3 × 2",
         "opencode.db 源库只读 + myagentwatch.db 聚合库读写；单文件零运维；无需独立数据库进程"],
        ["容器化", "Docker + docker-compose",
         "一键部署；卷挂载实现配置/数据持久化；:ro 只读挂载保护源数据"],
    ],
    col_widths=[2.2, 3.8, 8.5],
)

heading("3.2 为什么选择 SQLite 而不是 PostgreSQL？", level=2)
para(
    "MyAgentWatch 定位为轻量级自托管工具。SQLite 的优势在于：（1）零部署——不需要独立数据库进程，"
    "一个文件即完整数据库；（2）并发足够——AI Agent 团队通常 3-10 个 Agent，SQLite WAL 模式下轻松应对；"
    "（3）备份简单——cp 一个文件即完成备份；（4）嵌入式友好——未来可打包为单个二进制。"
    "当用户量或 Agent 数量扩展到 50+ 时，可平滑迁移至 PostgreSQL。"
)

heading("3.3 为什么需要两份 SQLite 数据库？", level=2)
para(
    "MyAgentWatch 采用「源库只读 + 聚合库读写」的双库设计。opencode.db 属于 OpenCode CLI 的内部数据，"
    "MyAgentWatch 仅以只读方式挂载，绝不修改源数据。所有聚合统计、告警记录、聊天消息、用户令牌等自有数据，"
    "存入独立的 myagentwatch.db。这样既保证了源数据安全，又实现了职责分离——源库升级或更换不影响监控数据。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 四、系统架构
# ═══════════════════════════════════════════════════════════════
heading("四、系统架构设计", level=1)

heading("4.1 整体架构：四层模型", level=2)
para(
    "MyAgentWatch 采用经典的四层架构，从上到下依次为：\n"
    "第一层 · 展示层 —— 浏览器 SPA + CLI 终端，负责数据可视化和用户交互\n"
    "第二层 · API 层 —— Flask REST API + WebSocket + SSE，负责请求路由和实时推送\n"
    "第三层 · 业务层 —— 采集调度器 + 状态机 + 告警引擎 + 事件总线 + 聊天系统，核心业务逻辑\n"
    "第四层 · 数据层 —— 源数据库（只读）+ 聚合数据库（读写）+ 日志归档，数据持久化"
)

heading("4.2 数据采集：适配器模式", level=2)
para(
    "MyAgentWatch 的核心设计模式是适配器模式（Adapter Pattern）。所有数据源实现统一的 SourceInterface 抽象基类，"
    "通过 @register_source 装饰器注册到 SOURCE_REGISTRY。新增数据源只需在 config.yaml 中添加一行配置，"
    "Collector 自动发现并实例化对应的适配器，无需修改一行核心代码。"
)

para("SourceInterface 要求每个适配器实现三个方法：", bold=True)
bullet("● discover_agents() → List[Agent]  —— 自动扫描数据源，发现新 Agent")
bullet("● collect(since_timestamp) → CollectedData  —— 增量采集，返回自上次同步以来的新数据")
bullet("● health_check() → Dict  —— 检查数据源自身是否健康（连接正常、数据可读）")

para("\n目前已实现的数据源适配器：", bold=True)

table(
    ["适配器", "数据源类型", "采集内容", "增量策略"],
    [
        ["OpenCodeDBSource", "SQLite 数据库",
         "session 表（会话树、标题、目录）、message 表（角色、模型、Token、成本）、part 表（thinking/tool/text/step 内容块）",
         "记录 last_sync_time，WHERE time_updated > last_sync"],
        ["OpenCodeLogSource", "日志文件（~/.local/share/opencode/log/*.log）",
         "LLM 调用耗时（elapsed）、服务启动/停止事件、工具注册清单、异常堆栈、权限评估日志",
         "记录文件名 + 行号偏移量，seek 到断点继续读"],
        ["ClaudeCodeSource", "Claude Code 日志文件",
         "Claude Code CLI 的原生日志格式，包括思考过程、工具调用、handoff 事件、文件读写操作",
         "同日志文件增量策略"],
        ["SystemSource", "psutil 系统调用",
         "CPU 使用率（每核 + 总体）、内存使用量/百分比、磁盘 I/O、磁盘使用率、网络 I/O",
         "每次全量采集（psutil 调用开销 < 1ms）"],
    ],
    col_widths=[2.5, 2.5, 6, 3.5],
)

heading("4.3 Agent 状态机：5 状态 + 精准判定", level=2)
para(
    "MyAgentWatch 实现了业界最精细的 Agent 状态模型（6 种状态），远超 Multica 的简单 4 状态模型。"
    "状态判定综合考虑三个信号源：主动心跳（heartbeat）、活动日志（activity_log）、源数据发现时间戳（last_seen_time），"
    "取 max(三者) 作为 actual_last，避免单一信号源失效导致误判。"
)

table(
    ["状态", "颜色", "图标", "含义", "触发条件"],
    [
        ["active", "绿色", "🟢", "正常运行中",
         "最近有心跳或活动记录 AND 无近期错误 AND 模型ID已配置"],
        ["working", "蓝色", "🔵", "正在执行任务",
         "Agent 显式上报 working 状态（通常在处理用户请求时自动切换）"],
        ["idle", "黄色", "🟡", "空闲或状态不确定",
         "新注册的 Agent 等待首次心跳 OR 活动时间戳超过阈值但非深度超时 OR 数据来源非心跳且不确定"],
        ["error", "红色", "🔴", "发生错误",
         "近期有 error/critical 级别的活动记录 OR 模型 ID 缺失（需模型却未配置）"],
        ["blocked", "橙色", "🟠", "被阻塞需人工介入",
         "error 状态持续超过 2 倍超时阈值（默认 600s），判定为持久性故障"],
        ["offline", "紫色", "🟣", "离线/不可达",
         "数据源被禁用/删除/断开 OR 心跳丢失且 heartbeat 为最新信号 OR 从未启动"],
    ],
    col_widths=[1.5, 1.2, 1, 2.5, 8.3],
)

para("\n状态机关键设计决策：", bold=True)
bullet("● 心跳优先 —— 有心跳的 Agent 使用显式心跳时间；无心跳的旧版 Agent 回退到 activity_log")
bullet("● offline 严格判定 —— 仅当 heartbeat 是最新信号时才将深度超时判为 offline，避免数据源发现时间戳干扰")
bullet("● 永不删除 —— Agent 不会因为离线被删除，只有从所有数据源中移除才会消失")
bullet("● 状态追溯 —— status_since 字段记录进入当前状态的时间戳，metadata 中存储 status_reason")

heading("4.4 事件驱动：双通道实时推送", level=2)
para(
    "MyAgentWatch 设计了双通道推送架构，区分结构化数据和非结构化事件流：\n\n"
    "WebSocket 通道（Flask-SocketIO）：\n"
    "• 承载结构化数据：Agent 状态变更（agent_update）、仪表盘统计卡片（stat_snapshot）、流程图增量（flow_update）、告警事件（alert_event）\n"
    "• 推送策略：增量更新 2 秒一次（仅推送变化的数据），全量快照 10 秒一次（保证最终一致性），新 WebSocket 连接即时推送当前快照\n"
    "• 房间机制：每个数据源对应一个 SocketIO room，客户端按需订阅\n\n"
    "SSE 通道（Server-Sent Events）：\n"
    "• 承载文本内容流：thinking（思考过程）、tool_call（工具调用详情）、response（输出文本）、handoff（Agent 交接）\n"
    "• 单向推送：服务端 → 客户端，无需客户端回复，天然适合事件展示\n"
    "• 自动重连：浏览器原生 EventSource API 内置断线重连机制\n\n"
    "关键优化：socketio.emit 调用在后台线程池中执行，不阻塞 Flask HTTP 响应主线程。"
    "避免了一个慢客户端拖慢整个系统的经典问题。"
)

heading("4.5 项目完整目录结构", level=2)

code("""myagentwatch/                          # ← 主项目根目录
├── app.py                              # Flask + SocketIO 应用入口
├── config.yaml                         # 用户配置文件（数据源、告警规则、模板选择）
├── requirements.txt                    # Python 依赖清单
├── check.py                            # 健康检查脚本
├── Dockerfile                          # Docker 镜像构建
├── docker-compose.yml                  # 一键部署编排
├── MyAgentWatch-架构方案.md             # 架构设计文档
├── MyAgentWatch-2.0-架构方案.md         # 2.0 升级方案
│
├── myagentwatch/                       # Python 包（核心业务逻辑）
│   ├── db.py                           # 数据库管理：16 张表建表、迁移、连接池
│   ├── config.py                       # 配置加载：YAML 解析 + 默认值 + 模板引擎 + 路径展开
│   ├── collector.py                    # 采集调度器（866 行核心文件）
│   │   · Collector 类：数据源初始化、collect_all() 主循环
│   │   · _mark_stale_agents()：6 状态机判定（3 信号源综合）
│   │   · _persist_data()：message/part/tool_call 三表并行写入
│   │   · _persist_turns()：对话 Turn 去重批量写入（executemany）
│   │   · _aggregate_daily_stats()：日聚合（每天一次，非每 2s）
│   │   · _publish_events()：Part 级事件发布到 SSE EventBus
│   │   · _archive_and_cleanup()：gzip 归档 + 超期删除 + VACUUM
│   │   · _build_relationships()：会话 parent→child 构建 Agent 关系图
│   ├── pricing.py                      # 定价模型：8 厂商 37 模型（含 DeepSeek/Qwen/Kimi/GLM/豆包）
│   ├── user.py                         # 用户管理：PAT 令牌生成/验证、SHA-256 哈希、myaw_ 前缀
│   ├── alerting.py                     # 告警规则引擎：config 驱动、阈值比较、alert 持久化
│   ├── event_bus.py                    # SSE 事件总线：publish/subscribe、分类过滤
│   ├── queries.py                      # 聚合查询：dashboard/by-agent/by-hour/by-model
│   ├── log_compiler.py                 # 对话日志编译器：Turn → ContentBlock 结构化
│   │
│   ├── sources/                        # 数据源适配器包
│   │   ├── __init__.py                 # SOURCE_REGISTRY + @register_source 装饰器
│   │   ├── base.py                     # SourceInterface 抽象基类 + CollectedData 数据类
│   │   ├── opencode_db.py              # OpenCode SQLite 适配器
│   │   ├── opencode_log.py             # OpenCode 日志文件适配器
│   │   ├── claude_code.py              # Claude Code 日志适配器（text block 补丁）
│   │   ├── sqlite_agent.py             # 通用 SQLite Agent 适配器
│   │   ├── log_adapter.py              # LogAdapter 接口（parse_turns）
│   │   └── system.py                   # 系统资源适配器（psutil）
│   │
│   ├── routes/                         # Flask Blueprint 路由
│   │   ├── api.py                      # REST API：40+ 端点（agents/sessions/stats/tokens/timeline/alerts/pricing...）
│   │   └── chat_api.py                 # 聊天 API：消息 CRUD、好友请求、任务分享、_emit_async 后台推送
│   │
│   └── templates/                      # 行业模板
│       ├── template_engine.py          # 模板加载 + 深度合并引擎
│       ├── default.yaml                # 默认模板
│       ├── quant_trading.yaml          # 量化交易行业模板
│       └── web_dev.yaml                # Web 开发行业模板
│
├── static/                             # 前端静态资源（SPA 单页应用）
│   ├── index.html                      # 入口 HTML（深色主题）
│   ├── css/
│   │   ├── dashboard.css               # 仪表盘样式（#1a1a2e 深色背景 + #16213e 卡片）
│   │   ├── chat-wechat.css             # 群聊 + 事件流 + Token + 收件箱 全部样式
│   │   └── log-viewer.css              # 日志查看器样式（彩色编码行）
│   └── js/
│       ├── app.js                      # 主入口：路由管理、全局状态、WebSocket 连接
│       ├── chat-wechat.js              # 微信风格三栏群聊（消息去重 + 富媒体气泡）
│       ├── event-stream.js             # SSE 事件流（Agent 筛选 + 类别 checkbox + 星标 + 展开详情）
│       ├── token-dashboard.js          # Token 仪表盘（柱状图 + 双表 + 趋势图 + 未映射诊断）
│       ├── topology.js                 # dagre + d3.js DAG 交互拓扑图
│       ├── tree-view.js                # 会话树形视图（展开/折叠 + 点击详情）
│       ├── tree-layout.js              # 树形布局算法
│       ├── tree-list-view.js           # 树形列表混合视图
│       ├── list-view.js                # 扁平列表视图
│       ├── log-viewer.js               # 日志查看器（过滤 + 搜索 + 彩色编码）
│       ├── node-detail.js              # 节点详情抽屉面板
│       ├── charts.js                   # Chart.js 图表封装（柱状图/折线图/趋势图）
│       ├── constants.js                # 常量定义（颜色映射、状态枚举、事件类型）
│       ├── util.js                     # 工具函数（时间格式化、数字缩写、防抖节流）
│       └── toast.js                    # Toast 通知组件（成功/错误/警告/信息）
│
├── templates/                          # 独立模板目录
│   ├── default.yaml
│   ├── quant_trading.yaml
│   └── web_dev.yaml
│
├── changelog/                          # 开发日志（按日记录）
│   ├── 2026-05-15.md ~ 2026-06-03.md
│
└── data/                               # 运行时数据（自动创建）
    ├── myagentwatch.db                 # 聚合数据库
    └── archive/                        # 日志归档（gzip JSONL）""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 五、核心功能详解
# ═══════════════════════════════════════════════════════════════
heading("五、核心功能详解", level=1)

# 5.1 Token 仪表盘
heading("5.1 Token 用量仪表盘 —— 省钱的第一道防线", level=2)
para(
    "Token 成本是 AI Agent 团队最大的运营支出。一个 5 人团队使用 Claude Opus 做代码生成，"
    "每月可能产生 $200-$500 的 Token 费用。MyAgentWatch 的 Token 仪表盘让每一分钱都清楚可见。"
)
para("功能要点：", bold=True)
bullet("● 五级 Token 拆解 —— input（输入消耗）/ output（输出生成）/ reasoning（思考链消耗）/ cache_read（缓存命中读取）/ cache_write（缓存写入），比业界常见的三级拆解更精细")
bullet("● 五种查询维度 —— dashboard（总览）/ by-agent（按 Agent 聚合）/ by-model（按模型聚合）/ by-hour（按小时趋势）/ unmapped（未定价模型诊断，帮助发现配置遗漏）")
bullet("● 8 家厂商 37 个模型完整定价 —— 国际：OpenAI (GPT-4o/4.1/o3)、Anthropic (Claude Opus/Sonnet/Haiku)、Google (Gemini)；国内：DeepSeek (V3/V4/R1)、阿里 (Qwen)、月之暗面 (Kimi)、智谱 (GLM)、字节 (豆包)")
bullet("● 前端三视图 —— 柱状图（按 Agent 和模型对比）、双数据表（详细用量 + 成本排序）、趋势折线图（按小时查看消耗波动）、未映射诊断面板（高亮显示缺少定价的模型）")
bullet("● 日聚合预计算 —— daily_stats 表每天聚合一次（非实时计算），避免每次 API 请求扫描全量 token_records")

# 5.2 事件流
heading("5.2 SSE 实时事件流 —— Agent 的「直播画面」", level=2)
para(
    "事件流是 MyAgentWatch 最具特色的功能之一。它将 Agent 的内部工作过程以彩色编码的实时流形式展现，"
    "让开发者像看直播一样观察 Agent 的思考、决策和执行。"
)
para("功能要点：", bold=True)
bullet("● 四种事件类型 —— thinking（蓝色·Agent 的思考推理过程）、tool_call（橙色·工具调用的名称/参数/耗时/退出码）、response（绿色·Agent 的文本输出）、handoff（紫色·子 Agent 交接事件）")
bullet("● 自动启动 —— 页面加载即连接 SSE（DOMContentLoaded 事件触发），不等用户切换到事件流 tab。tab 切换不断开连接，页面关闭才释放。这是 5 月 30 日根据用户反馈修复的关键 UX 问题")
bullet("● 强大的筛选能力 —— 按 Agent 分组下拉选择 + 按事件类别 checkbox 勾选 + 星标收藏 + 「仅显示错误」一键过滤")
bullet("● 详情展开 —— 点击任意事件行，下方展开详情面板，以 <pre> 标签展示完整文本内容（thinking 可能数千字），支持滚动查看")
bullet("● 操控面板 —— 暂停/恢复推送 + 一键清空列表")
bullet("● 技术实现 —— 后端 EventBus (event_bus.py) 在 collector 每轮采集后批量 publish Part 级事件；前端 EventSource API 接收，每个事件分配唯一 ID 用于去重")

# 5.3 群聊
heading("5.3 Agent 群聊 —— 微信风格的 Agent 社交系统", level=2)
para(
    "6 月 1 日全新重做的群聊系统，采用微信桌面版的三栏布局，让 Agent 之间的沟通像微信一样自然。"
    "这是 MyAgentWatch 相比 Multica 等竞品最重要的差异化功能——Multica 完全没有 Agent 社交能力。"
)
para("三栏布局详解：", bold=True)
bullet("● 左栏 · 会话列表（宽度 ~280px）—— 圆形头像 + 在线绿点 + 最后一条消息预览（截断） + 红点未读计数 + 相对时间戳（刚刚/5分钟前/昨天）")
bullet("● 中栏 · 消息流（弹性宽度）—— 人类消息：蓝色气泡靠右对齐 / Agent 消息：灰色气泡靠左对齐。富媒体消息类型：tool_call 橙色边框卡片（显示工具名+参数摘要）、handoff 紫色边框卡片（显示交接双方+子Agent类型）、share 绿色边框卡片（显示任务标题+结果摘要）")
bullet("● 右栏 · Agent 详情（宽度 ~300px）—— 上方：选中的 Agent 信息卡片（头像 + 显示名称 + 模型ID + 数据来源 + Token 消耗统计）。下方：通讯录列表（按 group_name 分组 + 在线绿点；active/idle/working 三种状态均视为在线）")
para("\n技术要点：", bold=True)
bullet("● 消息去重 —— 基于 natural_key 在数据库层面 UNIQUE 约束 + 内存 set 双重去重")
bullet("● 实时推送 —— chat_api 的 _emit_async 函数将 socketio.emit 提交到 ThreadPoolExecutor，不阻塞 HTTP 响应")
bullet("● 旧代码清零 —— 重做时完全移除旧版聊天代码，不留兼容包袱")

# 5.4 收件箱
heading("5.4 通知收件箱 —— 不遗漏任何重要信息", level=2)
para(
    "顶部导航栏铃铛图标 + 红色未读数字徽章。系统自动生成三类通知并写入 inbox 表："
)
bullet("● Agent 消息 —— 当有 Agent 在群聊中 @你或向你发送消息时触发")
bullet("● 好友请求 —— 当其他 Agent 发起好友申请时触发")
bullet("● 告警触发 —— 当告警规则被触发（如 tool_failure_rate > 20%）时自动生成通知")
bullet("● 一键全部标记已读 + 单条标记已读")

# 5.5 告警引擎
heading("5.5 告警规则引擎 —— 主动发现问题", level=2)
para(
    "内置 4 条默认告警规则，全部通过 config.yaml 的 alert_rules 数组配置，支持热修改。"
    "告警触发后写入 alerts 表，同时在 WebSocket 推送 alert_event，前端可弹出 Toast 通知。"
)

table(
    ["规则名称", "监控指标", "运算符", "阈值", "严重级别", "业务含义"],
    [
        ["agent_idle", "当前时间 - 最后活动时间", ">", "3600 秒", "warn",
         "Agent 超过 1 小时无任何活动，可能已停止工作或任务卡住"],
        ["high_cost", "单会话累计 Token 成本", ">", "$5.00", "warn",
         "单次对话花费超过 5 美元，需关注是否正常（大量分析任务）还是异常（死循环）"],
        ["tool_failure_rate", "失败工具调用 / 总工具调用 × 100%", ">", "20%", "critical",
         "工具调用失败率超过 20%，可能权限问题、API 变更、文件系统满等严重故障"],
        ["cache_hit_rate_low", "cache_read / (cache_read + cache_write) × 100%", "<", "30%", "info",
         "缓存命中率过低，可能 prompt 设计不佳导致每次请求都需重新计算，增加延迟和成本"],
    ],
    col_widths=[2.2, 3.5, 1, 1.5, 1.5, 5.3],
)

# 5.6 PAT 令牌
heading("5.6 PAT 令牌认证 —— 安全的 Agent 身份系统", level=2)
para(
    "MyAgentWatch 实现了完整的 PAT (Personal Access Token) 身份认证系统，用于 CLI 客户端和服务间通信。"
)
bullet("● 令牌格式 —— myaw_ 前缀（便于识别和审计）+ 40 位十六进制随机数")
bullet("● 存储安全 —— 数据库中仅存储 SHA-256 哈希值，明文令牌仅在生成时返回一次")
bullet("● 身份绑定 —— 每个令牌绑定到 specific agent_id，connect 时自动匹配 token_prefix 识别 Agent 名称")
bullet("● API 认证 —— 所有 /api/* 端点通过 Bearer Token 认证，前端页面免认证")
bullet("● 用户管理 —— 完整的 CRUD 端点：创建用户、生成令牌、列出用户、撤销令牌、删除用户")

# 5.7 性能优化
heading("5.7 性能优化历程（5 月 30 日）", level=2)
para(
    "5 月 30 日进行了一次集中的性能审计和优化，以下 5 项改进将系统负载降低了约 80%："
)
table(
    ["优化项", "优化前", "优化后", "效果"],
    [
        ["daily_stats 聚合频率", "每次采集循环（每 2 秒）", "每天一次（首次访问时触发）",
         "数据库写入减少 43,200 倍"],
        ["轮询间隔", "2 秒", "5 秒", "数据库读取压力降低 60%"],
        ["工作线程数", "4 线程", "8 线程", "采集 + 推送并行度翻倍"],
        ["socketio.emit 执行", "同步阻塞在 HTTP 响应线程", "提交到 ThreadPoolExecutor 后台执行",
         "HTTP 响应延迟从 200ms+ 降至 < 20ms"],
        ["actual_last 计算", "仅使用 heartbeat", "max(heartbeat, activity_log, last_seen_time)",
         "消除误判，旧版 Agent 不会因无心跳而误报离线"],
    ],
    col_widths=[3, 3, 3.5, 5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 六、数据模型
# ═══════════════════════════════════════════════════════════════
heading("六、数据模型设计", level=1)

heading("6.1 数据库概览", level=2)
para(
    "myagentwatch.db 包含 16 张核心数据表，按功能分为 5 个组：\n"
    "• 基础配置组：data_sources（数据源）、pricing（定价）、users（用户/令牌）\n"
    "• Agent 运行组：agents（Agent 注册）、agent_relationships（Agent 间调用关系）\n"
    "• 业务数据组：sessions（会话）、token_records（Token 记录）、tool_calls（工具调用）、activity_log（行为日志）\n"
    "• 对话分析组：conversation_turns（对话 Turn）、turn_content（Turn 内容块）、agent_handoffs（Agent 交接）\n"
    "• 运维管理组：daily_stats（日聚合）、alerts（告警）、health_checks（系统健康）、inbox（收件箱）"
)

heading("6.2 核心表设计说明", level=2)

para("agents 表 —— Agent 生命周期管理", bold=True)
bullet("● 复合主键：agent_id = source_name + agent_name + model_id，支持同一 Agent 跨数据源区分")
bullet("● 配置增强：discover_agents() 自动发现后，与 config.yaml 的 agent_meta 合并 → display_name + group_name")
bullet("● 状态字段：status（当前状态）+ status_since（进入时间）+ metadata.status_reason（原因）")

para("\ntoken_records 表 —— Token 消费的完整审计日志", bold=True)
bullet("● message 粒度：每条 message 生成一条记录（不是 part 粒度，因为 Token 数据在 message 级别）")
bullet("● 5 级 Token 字段：input / output / reasoning / cache_read / cache_write")
bullet("● 支持按 session_id / agent_id / model_id / timestamp 多维查询")

para("\nconversation_turns + turn_content —— 对话结构化存储", bold=True)
bullet("● natural_key：业务主键（agent_id + session_id + seq），用于去重")
bullet("● 一对多关系：一个 Turn 可包含多个 ContentBlock（如思考块 + 工具调用块 + 输出文本块）")
bullet("● handoff_id：外键关联 agent_handoffs，追踪子 Agent 委托链")

para("\nagent_handoffs 表 —— Agent 委托链追踪", bold=True)
bullet("● 记录完整交接信息：from_agent_id → to_agent_id，委托 prompt，返回 result")
bullet("● subagent_type：区分 Agent / Task / Subagent 不同类型")
bullet("● to_session_id：交接后在子 Agent 侧创建的新会话 ID")

heading("6.3 定价模型表", level=2)
para(
    "pricing 表存储每个模型每百万 Token 的价格（美元）。支持字段：provider（厂商）、model_id（模型标识）、"
    "input_price（输入价格/百万Token）、output_price（输出价格/百万Token）、cache_price（缓存读写价格）。"
    "成本计算公式：cost = (tokens_input × input_price + tokens_output × output_price + tokens_reasoning × reasoning_price + cache_read × cache_price + cache_write × cache_price) / 1,000,000。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 七、API 设计
# ═══════════════════════════════════════════════════════════════
heading("七、API 设计", level=1)

heading("7.1 REST API（40+ 端点）", level=2)
para("MyAgentWatch 的 REST API 按资源分组，提供 JSON 格式响应。所有 /api/* 端点需要 Bearer Token 认证。")

table(
    ["资源组", "端点", "方法", "说明"],
    [
        ["Agent", "/api/agents", "GET", "获取所有 Agent 列表（含实时状态、模型、所属数据源）"],
        ["Agent", "/api/agents/<id>", "GET", "获取单个 Agent 详情 + 最近活动记录 + Token 统计"],
        ["会话", "/api/sessions", "GET", "获取活跃/历史会话列表（支持 status/agent_id 过滤）"],
        ["会话", "/api/sessions/<id>", "GET", "获取会话详情（标题、目录、完整 message→part 时间线）"],
        ["统计", "/api/stats/overview", "GET", "仪表盘概览：活跃 Agent 数、今日 Token 总量、成功率、今日成本"],
        ["Token", "/api/stats/tokens", "GET", "Token 综合数据：by_day + by_model + overview"],
        ["Token", "/api/stats/tokens/by-agent", "GET", "按 Agent 聚合：每个 Agent 的 input/output/cost 总和"],
        ["Token", "/api/stats/tokens/by-hour", "GET", "按小时趋势：最近 24h 每小时的 Token 消耗曲线"],
        ["Token", "/api/stats/tokens/by-model", "GET", "按模型聚合：每种模型的用量和成本排序"],
        ["Token", "/api/stats/tokens/unmapped", "GET", "未映射诊断：pricing 表中没有的模型及其 Token 量"],
        ["图表", "/api/stats/charts", "GET", "Chart.js 数据：延迟折线 + Token 柱状图原始数据"],
        ["时间线", "/api/timeline", "GET", "活动时间线：按时间排序的 message + tool_call 混合流"],
        ["时间线", "/api/timeline/flow", "GET", "DAG 流程图数据：节点（session/message）+ 边（parent/child/tool_call）"],
        ["事件流", "/api/events/stream", "GET", "SSE 端点：Content-Type: text/event-stream 实时推送"],
        ["日志", "/api/logs", "GET", "实时日志流（SSE）：按 level/service 过滤的彩色日志"],
        ["健康", "/api/health", "GET", "系统运维状态：各数据源连接状态 + CPU/内存/磁盘指标"],
        ["告警", "/api/alerts", "GET", "告警列表（支持 active/resolved 过滤）"],
        ["告警", "/api/alerts/resolve", "POST", "手动解除指定告警"],
        ["配置", "/api/config", "GET", "当前运行配置快照（不含敏感信息）"],
        ["定价", "/api/pricing", "GET", "所有模型定价表"],
        ["聊天", "/api/chat/messages/<conv_id>", "GET/POST", "读取/发送群聊消息"],
        ["聊天", "/api/chat/friend-request", "POST", "发送好友请求"],
        ["聊天", "/api/chat/agent-message", "POST", "Agent 发布动态到群聊"],
        ["聊天", "/api/chat/share-task/<conv_id>", "POST", "分享任务成果到指定会话"],
        ["收件箱", "/api/inbox", "GET", "获取通知列表（支持 limit/unread 过滤）"],
        ["收件箱", "/api/inbox/<id>/read", "POST", "标记单条已读"],
        ["收件箱", "/api/inbox/read-all", "POST", "全部标记已读"],
        ["用户", "/api/users", "GET/POST", "列出/创建用户"],
        ["用户", "/api/users/<id>/token", "POST", "为用户生成新的 PAT 令牌"],
        ["心跳", "/api/heartbeat/<agent_id>", "POST", "接收 Agent 主动心跳（更新状态和 last_heartbeat_at）"],
    ],
    col_widths=[1.5, 4, 1.2, 7.8],
)

heading("7.2 WebSocket 事件", level=2)

table(
    ["方向", "事件名", "触发时机", "负载内容"],
    [
        ["S→C", "agent_update", "Agent 状态变更时",
         "{agent_id, status, status_since, status_reason, tokens_1h, model_id}"],
        ["S→C", "stat_snapshot", "每 10 秒全量快照 + 新连接即时推送",
         "{active_agents, total_tokens, success_rate, today_cost, agents[]}"],
        ["S→C", "alert_event", "告警规则触发时",
         "{rule_name, level, message, agent_id, timestamp}"],
        ["S→C", "flow_update", "DAG 流程图有增量变化时",
         "{type: add/update/remove, node_id, parent_id, data}"],
        ["S→C", "chat_message", "新聊天消息时",
         "{conversation_id, sender_name, content, sender_type, timestamp}"],
        ["C→S", "subscribe_logs", "客户端请求订阅日志",
         "{source_id, filters: {level, service, keyword}}"],
        ["C→S", "unsubscribe_logs", "客户端取消日志订阅",
         "{source_id}"],
    ],
    col_widths=[1.2, 2.8, 4, 6.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 八、CLI 客户端
# ═══════════════════════════════════════════════════════════════
heading("八、CLI 客户端（myagentwatch-cli）", level=1)

heading("8.1 设计理念", level=2)
para(
    "myagentwatch-cli 是一个独立的命令行客户端程序，与主项目分开维护、独立打包。"
    "设计目标是为 Agent 本身提供一种轻量级的、可脚本化的方式与 MyAgentWatch 服务端交互。"
    "纯 Python 标准库实现（urllib.request），零外部依赖，确保在任何 Python 3.10+ 环境可直接运行。"
)

heading("8.2 项目结构", level=2)
code("""myagentwatch-cli/
├── pyproject.toml                  # Python 包配置
│   · name: myagentwatch-cli
│   · version: 0.1.0
│   · entry_point: myaw = myagentwatch_cli.cli:main
│   · requires-python: >=3.10
│   · build-system: setuptools
├── config.json                     # 运行时配置（自动生成）
│   · server: 服务端地址（如 http://127.0.0.1:10000）
│   · key: PAT 令牌（如 myaw_ad7e21...）
│   · agent_name: 自动识别的 Agent 名称
│   · agent_id: 自动识别的 Agent ID
└── myagentwatch_cli/
    ├── __init__.py
    ├── cli.py                      # 命令行入口（argparse 子命令 + 11 个命令实现）
    └── client.py                   # HTTP 客户端（GET/POST/DELETE + Bearer 认证）""")

heading("8.3 client.py —— HTTP 通信层详解", level=2)
para(
    "client.py 是整个 CLI 的网络基础层，负责所有与 MyAgentWatch 服务端的 HTTP 通信。"
    "全部使用标准库 urllib.request 实现，不依赖 requests/httpx 等第三方库。"
)
bullet("● Bearer Token 认证 —— 所有请求自动添加 Authorization: Bearer <token> 请求头")
bullet("● URL 安全编码 —— urllib.parse.quote(path, safe=\"/?=&:%\")，处理 Agent ID 中的空格和冒号")
bullet("● config.json 持久化 —— load_config() / save_config() 读写 JSON 配置文件，存储 server + key + agent_name + agent_id")
bullet("● connect() 自动身份识别 —— 连接后调用 GET /api/users，遍历用户列表匹配 token_prefix 找到对应的 agent_name 和 agent_id，更新 config.json")
bullet("● 超时和错误处理 —— 默认 30s 超时；HTTP 错误返回 JSON 解析结果；网络异常返回 {\"error\": str(e)}")
bullet("● 三个 HTTP 方法封装 —— get(path) / post(path, body) / delete(path)")

heading("8.4 cli.py —— 命令行接口详解", level=2)
para(
    "cli.py 使用 Python 标准库 argparse 实现子命令模式。入口点 myaw（由 pyproject.toml 的 [project.scripts] 定义）。"
    "终端输出使用 ANSI 转义序列实现彩色文本和 Unicode 边框卡片。"
)
bullet("● ANSI 颜色映射 —— active=32(绿) / working=34(蓝) / idle=33(黄) / error=31(红) / blocked=35(紫) / offline=37(灰)")
bullet("● _box() 函数 —— 用 Unicode 字符绘制带标题的边框卡片（┌─┐│└─┘），自适应内容宽度")
bullet("● 每个命令对应一个 cmd_*() 函数，调用 client.py 的 get/post 方法，格式化输出")

heading("8.5 11 个命令完整说明", level=2)

table(
    ["命令", "参数", "API 调用", "输出内容"],
    [
        ["connect", "--server --key",
         "GET /api/users",
         "连接服务端 → 保存 config.json → 识别 Agent 身份 → 显示「已连接」卡片（服务端地址+版本+运行时长）"],
        ["status", "无",
         "GET /api/agents + GET /api/tokens/dashboard + GET /api/inbox",
         "仪表盘总览卡片：Agent 列表（状态色 + 名称 + 模型ID）+ Token Top5（按模型，tokens 量 + 成本 $）+ 未读通知数"],
        ["dashboard", "无",
         "status 的所有 API + feed 的 API",
         "status 输出 + 换行 + feed 输出（合并视图）"],
        ["agents", "无",
         "GET /api/agents + GET /api/users",
         "Agent 列表：状态色 + 名称 + 模型ID + 🔑 令牌标识（有关联 PAT 的 Agent 显示钥匙图标）"],
        ["chat", "[message] [--conv N]",
         "有 message: POST /api/chat/messages/N | 无 message: GET /api/chat/messages/N?limit=20",
         "发送模式：确认「已发送」+ 内容预览 | 读取模式：最近 20 条消息（时间 + 角色图标 + 发送者 + 内容，倒序显示）"],
        ["heartbeat", "--agent-id [--status] [--daemon]",
         "POST /api/heartbeat/{agent_id}",
         "单次模式：发送一次心跳，显示返回状态 | 守护模式：每 15s 循环发送，按 Ctrl+C 停止（用于无内置心跳的 Agent CLI）"],
        ["tokens", "[--days N]",
         "GET /api/tokens/dashboard + /api/tokens/by-agent + /api/tokens/unmapped",
         "Token 用量卡片：按日汇总 + 按 Agent Top5（名称 + Token 量 + 成本 + 任务数）+ 未定价模型警告（模型名 + Token 量）"],
        ["feed", "无",
         "GET /api/inbox?limit=30",
         "动态流列表：每条显示未读标识（●/○）+ 类型图标（💬消息/🤝好友/🚨告警/📤分享）+ 时间 + 标题 + 摘要（截断 100 字）"],
        ["post", "<content>",
         "POST /api/chat/agent-message",
         "发布 Agent 动态到群聊 → 确认「已发布」"],
        ["friend", "<agent_id> [message]",
         "POST /api/chat/friend-request",
         "发送好友请求 → 确认「已发送 → agent_id」"],
        ["share", "<title> [summary]",
         "POST /api/chat/share-task/1",
         "分享任务成果到群聊 → 确认「已分享：标题」"],
    ],
    col_widths=[1.5, 3, 4, 6],
)

heading("8.6 心跳守护模式详解", level=2)
para(
    "守护模式（--daemon）是 CLI 的关键功能。并非所有 Agent CLI 都内置了向 MyAgentWatch 发送心跳的能力。"
    "通过启动一个后台守护进程，CLI 客户端可以代理这些 Agent 维持 active 状态。"
)
bullet("● 运行方式 —— myaw heartbeat --agent-id \"...\" --daemon")
bullet("● 心跳间隔 —— 固定 15 秒（在 while True 循环中 time.sleep(15)）")
bullet("● 停止方式 —— Ctrl+C（KeyboardInterrupt 捕获，优雅退出）")
bullet("● 适用场景 —— 监控不支持主动心跳的旧版 Agent CLI、监控第三方 Agent 工具、临时为某个 Agent 注入心跳")

heading("8.7 安装与启动", level=2)
code("""# 方式一：从源码安装（推荐）
cd myagentwatch-cli
pip install -e .

# 方式二：直接运行
python -m myagentwatch_cli.cli status

# 首次使用：连接服务端
myaw connect --server http://localhost:10000 --key myaw_ad7e21...

# 日常使用
myaw status          # 查看仪表盘
myaw tokens --days 7 # 查看本周 Token
myaw chat "任务完成"  # 发送消息""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 九、前端设计
# ═══════════════════════════════════════════════════════════════
heading("九、前端设计", level=1)

heading("9.1 设计原则", level=2)
bullet("● 零构建工具链 —— CDN 引入 Chart.js / dagre / d3.js / socket.io-client，无需 npm/webpack/vite")
bullet("● 深色主题优先 —— 面向开发者群体，深色背景减少视觉疲劳，长时间监控更舒适")
bullet("● 响应式优先 —— SPA 单页应用，支持 PC / 平板 / 手机三种断点")
bullet("● 自动启动 —— 数据采集和 SSE 连接在 DOMContentLoaded 即启动，不等用户操作")

heading("9.2 视觉设计规范", level=2)
table(
    ["设计元素", "颜色/数值", "用途"],
    [
        ["页面背景", "#1a1a2e（深蓝黑）", "主背景色，降低屏幕亮度"],
        ["卡片背景", "#16213e（深蓝）", "数据卡片、面板、表格行交替"],
        ["主强调色", "#0f3460（中蓝）", "标题、重点数据、按钮"],
        ["次强调色", "#e94560（红粉）", "告警、错误状态、关键高亮"],
        ["active 状态", "#2ecc71（翠绿）", "Agent 在线标识、成功提示"],
        ["idle 状态", "#f1c40f（琥珀黄）", "Agent 空闲标识"],
        ["error 状态", "#e74c3c（红）", "Agent 错误标识、删除按钮"],
        ["字体", "系统默认等宽 / 14px", "代码感，适合开发者"],
    ],
    col_widths=[2.5, 4, 8],
)

heading("9.3 8 个功能页面", level=2)

table(
    ["页面", "路由", "核心功能"],
    [
        ["仪表盘", "#dashboard",
         "4 张概览卡片（活跃 Agent 数 / 今日 Token / 成功率 / 今日成本）+ Agent Token 柱状图（Chart.js bar chart）+ 1 小时请求延迟折线图 + Agent 实时状态表格（可排序）"],
        ["拓扑图", "#topology",
         "dagre 自动布局 + d3.js 渲染的 DAG 交互流程图。节点 = session/message，边 = parent→child 关系。支持拖拽、缩放、点击节点弹出详情抽屉（显示完整 message/part 内容）"],
        ["树形视图", "#tree",
         "基于会话 parent_id 构建的树形结构。支持展开/折叠所有节点、按 Agent 分组、点击节点弹出详情面板"],
        ["日志查看器", "#logs",
         "WebSocket 实时推送的彩色编码日志流。INFO=灰色 / WARN=黄色 / ERROR=红色 / DEBUG=浅蓝。支持按 level 和 service 过滤、关键词搜索"],
        ["群聊", "#chat",
         "微信风格三栏布局（见 5.3 节详细说明）。WebSocket 实时推送新消息，消息去重，富媒体气泡"],
        ["事件流", "#events",
         "SSE 实时推送（见 5.2 节详细说明）。分组 Agent 下拉 + 类别 checkbox + 星标 + 仅错误 + 暂停/清空 + 详情展开"],
        ["Token 仪表盘", "#tokens",
         "柱状图 + 按 Agent 数据表 + 按模型数据表 + 按小时趋势折线图 + 未映射模型诊断面板（红色高亮）"],
        ["收件箱", "#inbox",
         "通知列表（按时间倒序）+ 未读高亮 + 类型图标 + 全部标记已读按钮"],
    ],
    col_widths=[2, 2, 10.5],
)

heading("9.4 前端技术要点", level=2)
bullet("● 路由管理 —— 基于 window.location.hash 的轻量路由，不引入 Vue Router/React Router")
bullet("● 状态管理 —— 全局 appState 对象 + 事件触发式更新，不引入 Redux/Vuex")
bullet("● WebSocket 重连 —— socket.io 内置指数退避重连（1s → 2s → 4s → ... → 30s 上限）")
bullet("● SSE 重连 —— EventSource 浏览器原生自动重连（3s 间隔）")
bullet("● 消息去重 —— 两层去重：数据库 UNIQUE(natural_key) + 前端 Set<msg_id> 内存去重")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十、与 Multica 的竞品对比
# ═══════════════════════════════════════════════════════════════
heading("十、竞品对比：MyAgentWatch vs Multica", level=1)

heading("10.1 Multica 简介", level=2)
para(
    "Multica 是目前市场上功能最完整的 Agent 管理平台之一。采用 Go 1.26 + PostgreSQL 17 + Redis + WebSocket + "
    "Next.js 16 技术栈，支持 11 种 Agent CLI（Claude Code、Codex、Copilot、Cursor、Gemini、Kimi 等）。"
    "架构分为 5 层：Agent Backend 抽象层 → 本地守护进程 → HTTP 处理程序 → 领域服务 → 实时系统。"
    "核心优势在于 Token 用量分析（多维切片）和完整的任务生命周期管理（queued→dispatched→running→completed/failed/cancelled）。"
)

heading("10.2 功能对比矩阵", level=2)

table(
    ["能力维度", "MyAgentWatch", "Multica", "差异化分析"],
    [
        ["Agent 状态监控", "✅ 6 状态机（active/working/idle/error/blocked/offline）",
         "✅ 5 状态（idle/working/blocked/error/offline）",
         "MyAgentWatch 多了 working 状态（显式上报），区分更精细"],
        ["Token 用量分析", "✅ 8 厂商 37 模型 / 5 级拆解 / 5 种查询维度",
         "✅ 多维切片（按日/时/模型/Agent）",
         "MyAgentWatch 模型覆盖更广（尤其国内厂商），Token 拆解更细"],
        ["系统资源监控", "✅ CPU / 内存 / 磁盘 / 网络（psutil）",
         "❌ 不支持",
         "MyAgentWatch 独有，运维场景刚需"],
        ["Agent 群聊社交", "✅ 微信风格三栏 / 富媒体气泡 / 通讯录 / 好友",
         "❌ 完全没有",
         "MyAgentWatch 核心差异化武器"],
        ["任务生命周期", "计划中（6/10 截止）",
         "✅ 完整状态流转 / 时长统计",
         "Multica 当前领先，MyAgentWatch 即将补齐"],
        ["告警系统", "✅ 内置规则引擎 / config 驱动 / 4 条默认规则",
         "❌ 不支持",
         "MyAgentWatch 独有，主动发现问题的能力"],
        ["SSE 事件流", "✅ 4 类事件 / 彩色编码 / 筛选 / 详情展开",
         "❌ 不支持",
         "MyAgentWatch 独有，Agent 工作过程可视化"],
        ["CLI 客户端", "✅ 11 个命令 / 心跳守护 / 纯标准库",
         "✅ CLI 连接 Agent",
         "两者都有，MyAgentWatch 更偏 Agent 自助管理"],
        ["技术栈复杂度", "Python + SQLite → 单文件数据库，pip install 即用",
         "Go + PostgreSQL + Redis → 需数据库和缓存中间件",
         "MyAgentWatch 部署成本极低，适合个人/小团队"],
        ["多 Agent CLI 支持", "Claude Code / OpenCode 及兼容工具",
         "11 种 CLI（Claude Code / Codex / Copilot / Cursor / Gemini / Kimi 等）",
         "Multica 覆盖面更广，MyAgentWatch 聚焦核心"],
        ["Agent 关系图", "✅ 基于会话 parent→child 自动构建",
         "✅ Agent 拓扑图",
         "两者都有，实现方式不同"],
        ["通知收件箱", "✅ inbox 表 + 铃铛 + 未读徽章 + 自动通知",
         "✅ inbox_item 表",
         "两者都有"],
    ],
    col_widths=[2.5, 4.5, 4, 3.5],
)

heading("10.3 MyAgentWatch 的差异化定位", level=2)
para(
    "MyAgentWatch 不与 Multica 正面竞争「任务编排」和「多 CLI 广度」。我们的核心差异化在于：\n\n"
    "1. Agent 社交（群聊 + 通讯录 + 好友 + 动态）—— Multica 完全没有的方向，为 Agent 之间的协作提供社交基础设施\n"
    "2. 运维监控（CPU/内存/磁盘 + 告警规则引擎）—— 填补 Multica 的系统资源监控空白\n"
    "3. 部署成本（SQLite 单文件 + pip install）—— 面向个人开发者和小团队，无需运维 PostgreSQL/Redis\n"
    "4. 事件流可视化（SSE 实时流 + 彩色编码 + 详情展开）—— 让 Agent 工作过程完全透明\n\n"
    "一句话总结：Multica 管任务，MyAgentWatch 管 Agent。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十一、技术亮点与设计决策
# ═══════════════════════════════════════════════════════════════
heading("十一、技术亮点与设计决策", level=1)

heading("11.1 适配器模式 + 插件注册", level=2)
para(
    "通过 @register_source 装饰器将数据源适配器注册到 SOURCE_REGISTRY 字典。"
    "在 config.yaml 中新增数据源只需一行 type 声明，Collector 自动查找注册表并实例化。"
    "新增数据源类型无需修改 collector.py、api.py、前端代码。真正做到「对扩展开放，对修改封闭」。"
)

heading("11.2 双通道推送架构", level=2)
para(
    "区分结构化数据（WebSocket）和非结构化事件流（SSE）是经过权衡的设计决策。"
    "WebSocket 适合双向小消息的增量同步，SSE 适合单向大文本的事件推送。"
    "两者互补而非竞争——如果只用 WebSocket 推大文本，会阻塞状态更新通道；如果只用 SSE 推结构数据，会失去双向交互能力。"
)

heading("11.3 去重策略", level=2)
bullet("● conversation_turns：natural_key UNIQUE（DB层）+ 内存 Set 预查（应用层），双重去重")
bullet("● 聊天消息：消息 ID Set 内存去重，WebSocket 重复推送的包直接丢弃")
bullet("● Token 记录：INSERT OR IGNORE + message_id 唯一约束")

heading("11.4 配置优先于硬编码", level=2)
para(
    "整个系统贯彻「配置驱动」原则：数据源列表、Agent 元数据、告警规则、行业模板、轮询间隔、归档策略——"
    "全部可通过 config.yaml 修改，无需改代码。模板引擎支持按行业（量化交易 / Web 开发）一键切换预设配置。"
)

heading("11.5 日志归档与自动清理", level=2)
para(
    "conversation_turns 数据量随时间线性增长。系统实现了自动生命周期管理：\n"
    "• log_archive_days（默认 7 天）：超过此天数的 Turn 导出为 gzip JSONL 归档（按年月分文件），从主表删除\n"
    "• log_retention_days（默认 365 天）：超过此天数的数据彻底删除\n"
    "• 归档/删除后自动 VACUUM，回收磁盘空间\n"
    "这使得 myagentwatch.db 的大小保持在可控范围，长期运行也不会磁盘告警。"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十二、路线图
# ═══════════════════════════════════════════════════════════════
heading("十二、路线图（2026 年 6 月）", level=1)

heading("12.1 短期目标（6 月 10 日截止）", level=2)

table(
    ["优先级", "任务", "工期", "交付物"],
    [
        ["P1 🔴", "Agent Backend 抽象层", "2-3 天",
         "统一 SourceInterface 的子类 AgentBackend，支持通过统一接口启动/停止/配置多种 Agent CLI。"
         "首批支持 Claude Code 和 OpenCode。"],
        ["P2 🟡", "基础任务系统", "5-7 天",
         "新增 tasks 表 + task_records 表，实现 queued→assigned→running→completed/failed 任务生命周期。"
         "提供 /api/tasks CRUD 端点 + 前端任务看板页面。"],
        ["P3 🟢", "Agent 微信强化", "3-4 天",
         "群聊增加：@提及、消息引用回复、表情回应、群公告、消息搜索。通讯录增加：在线状态实时更新、Agent 签名/标签。"],
    ],
    col_widths=[1.5, 3.5, 1.5, 8],
)

heading("12.2 中期规划（6-7 月）", level=2)
bullet("● 多语言前端 —— 支持中文/英文/日文界面切换")
bullet("● 移动端适配 —— PWA + 响应式优化，手机端完整操作体验")
bullet("● 插件市场 —— 第三方数据源适配器、告警规则模板、仪表盘 Widget")
bullet("● PostgreSQL 迁移工具 —— 为 50+ Agent 团队提供数据库升级路径")

heading("12.3 不做的事情", level=2)
para("明确排除以下方向，保持项目聚焦：")
bullet("● Squad/小队管理 —— Multica 已做好，不做重复建设")
bullet("● Autopilot 自动执行 —— 安全风险高，保持「监控」定位")
bullet("● Electron 桌面客户端 —— Web + CLI 已足够，Electron 打包太重")
bullet("● 多工作区管理 —— 可通过多个 MyAgentWatch 实例组合实现")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 十三、部署与运维
# ═══════════════════════════════════════════════════════════════
heading("十三、部署与运维", level=1)

heading("13.1 Docker 部署（推荐）", level=2)
para(
    "最简单的部署方式，适合有 Docker 环境的团队。"
    "通过卷挂载实现配置和数据持久化，源数据库以只读方式挂载保护数据安全。"
)
code("""# docker-compose.yml
services:
  myagentwatch:
    build: .
    ports:
      - "10000:10000"                              # 映射到宿主机 10000 端口
    volumes:
      - ~/.local/share/opencode:/data/opencode:ro  # 源数据库：只读
      - ./config.yaml:/app/config.yaml              # 配置文件
      - ./data:/app/data                            # 聚合数据库 + 归档（读写）
    restart: unless-stopped                         # 自动重启""")

heading("13.2 本地开发部署", level=2)
code("""# 1. 创建虚拟环境
cd myagentwatch
python -m venv .venv
.venv\\Scripts\\activate    # Windows
source .venv/bin/activate   # macOS/Linux

# 2. 安装依赖
pip install -r requirements.txt

# 3. 启动服务端（默认端口 10000）
python app.py
# 服务启动后访问 http://localhost:10000

# 4. 健康检查
python check.py""")

heading("13.3 CLI 客户端安装", level=2)
code("""# 安装 CLI
cd myagentwatch-cli
pip install -e .

# 连接到已运行的服务端
myaw connect --server http://localhost:10000 --key myaw_xxx

# 日常使用
myaw status        # 仪表盘
myaw tokens        # Token 用量
myaw agents        # Agent 列表
myaw chat "hello"  # 发送消息""")

heading("13.4 配置说明", level=2)
para(
    "config.yaml 是系统的核心配置文件。所有配置项都有合理默认值，无此文件时系统使用默认配置自动扫描默认路径。"
    "以下为关键配置项说明："
)
bullet("● data_sources —— 数据源列表，每项包含 name/type/db_path/log_dir/enabled。支持配置多个数据源同时监控")
bullet("● agent_meta —— 为自动发现的 Agent 设置 display_name（显示名称）和 group（分组），可空")
bullet("● alert_rules —— 告警规则数组，每项包含 name/metric/condition/threshold/level，支持热修改")
bullet("● poll_interval —— 采集间隔（秒），默认 5 秒")
bullet("● agent_stale_timeout —— Agent 超时阈值（秒），默认 300 秒，超过此值进入 idle 状态")
bullet("● log_archive_days —— 日志归档天数（默认 7），超过此天数的对话 Turn 导出为 gzip 归档")
bullet("● log_retention_days —— 日志保留天数（默认 365），超过此天数的数据彻底删除")
bullet("● template —— 行业模板选择（quant_trading / web_dev / default），模板会与用户配置深度合并")

heading("13.5 运维要点", level=2)
bullet("● 数据库备份 —— 直接 cp myagentwatch.db 即可，SQLite 单文件零依赖备份")
bullet("● 磁盘空间监控 —— 关注 data/archive/ 目录大小，定期清理或增大 log_retention_days")
bullet("● 端口冲突 —— 默认 10000 端口，可在 Docker 或 app.py 启动参数中修改")
bullet("● 多实例部署 —— 每个 OpenCode 实例对应一个 MyAgentWatch 实例，通过不同端口隔离")

add_faq_chapter(doc, heading, para, bullet, code, table, callout)



# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "MyAgentWatch-项目介绍-v7.docx")
doc.save(output_path)
print(f"文档已生成: {output_path}")
print("共 13 章，可直接用于豆包/PPT 生成")
